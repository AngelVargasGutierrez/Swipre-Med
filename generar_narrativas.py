"""
Genera NARRATIVAS-MOPGIMED.docx con las 25 narrativas de casos de uso.
Convierte el formato Markdown al mismo estilo de tabla en Word.
"""

import glob
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colores ────────────────────────────────────────────────────────────────
C_ROJO        = 'B91C1C'   # encabezado de tabla principal
C_ROJO_CLARO  = 'FEF2F2'   # filas alternas en tablas de 3 col
C_GRIS_CLARO  = 'F1F5F9'   # columna "Campo" en tabla info
C_BLANCO      = 'FFFFFF'
C_OSCURO      = '1E293B'


# ─── Utilidades de formato de celda ─────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def hex_to_rgb(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_rich_text(paragraph, text, bold_default=False,
                  font_size=10, font_color=None):
    """Agrega texto con soporte para **negrita** en Markdown."""
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after  = Pt(1)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        is_bold = part.startswith('**') and part.endswith('**')
        content = part[2:-2] if is_bold else part
        if content:
            run       = paragraph.add_run(content)
            run.bold  = is_bold or bold_default
            run.font.size = Pt(font_size)
            run.font.name = 'Calibri'
            if font_color:
                run.font.color.rgb = hex_to_rgb(font_color)


# ─── Parseo de tablas Markdown ───────────────────────────────────────────────
def parse_row(line):
    line = line.strip()
    if line.startswith('|'): line = line[1:]
    if line.endswith('|'):   line = line[:-1]
    return [c.strip() for c in line.split('|')]


def is_separator(cells):
    return all(re.match(r'^-+$', c.replace(' ', '')) for c in cells if c)


def collect_table(lines, start):
    """Devuelve (filas_datos, índice_siguiente_línea)."""
    rows = []
    i    = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        cells = parse_row(lines[i])
        if not is_separator(cells):
            rows.append(cells)
        i += 1
    return rows, i


# ─── Creación de tablas en Word ──────────────────────────────────────────────
def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        if col_idx < len(row.cells):
            row.cells[col_idx].width = Cm(width_cm)


def build_info_table(doc, rows):
    """Tabla de 2 columnas: Campo | Detalle."""
    if not rows:
        return
    n_cols = 2
    tbl    = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows):
        row      = tbl.rows[r_idx]
        is_head  = r_idx == 0
        for c_idx in range(n_cols):
            cell_txt = row_data[c_idx] if c_idx < len(row_data) else ''
            cell     = row.cells[c_idx]
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if is_head:
                set_cell_bg(cell, C_ROJO)
                add_rich_text(para, cell_txt, bold_default=True,
                              font_size=10, font_color=C_BLANCO)
            else:
                if c_idx == 0:
                    set_cell_bg(cell, C_GRIS_CLARO)
                    add_rich_text(para, cell_txt, bold_default=True, font_size=10)
                else:
                    add_rich_text(para, cell_txt, font_size=10)

    # Ancho de columnas
    for row in tbl.rows:
        if len(row.cells) >= 2:
            row.cells[0].width = Cm(4.2)
            row.cells[1].width = Cm(11.3)

    doc.add_paragraph()


def build_events_table(doc, rows):
    """Tabla de 3 columnas: # | Usuario | Sistema."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    n_cols = max(n_cols, 3)
    tbl    = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows):
        row     = tbl.rows[r_idx]
        is_head = r_idx == 0
        for c_idx in range(n_cols):
            cell_txt = row_data[c_idx] if c_idx < len(row_data) else ''
            cell     = row.cells[c_idx]
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if is_head:
                set_cell_bg(cell, C_ROJO)
                add_rich_text(para, cell_txt, bold_default=True,
                              font_size=10, font_color=C_BLANCO)
            else:
                if r_idx % 2 == 0:
                    set_cell_bg(cell, C_ROJO_CLARO)
                add_rich_text(para, cell_txt, font_size=10)

    # Ancho de columnas
    for row in tbl.rows:
        if len(row.cells) >= 3:
            row.cells[0].width = Cm(1.0)
            row.cells[1].width = Cm(7.25)
            row.cells[2].width = Cm(7.25)

    doc.add_paragraph()


# ─── Procesador de archivo Markdown ─────────────────────────────────────────
def process_file(doc, filepath, is_first):
    with open(filepath, encoding='utf-8') as f:
        lines = f.readlines()
    lines = [l.rstrip('\n') for l in lines]

    if not is_first:
        doc.add_page_break()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # H1
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.runs[0].font.name = 'Calibri'
            i += 1
            continue

        # H2
        if line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.runs[0].font.name = 'Calibri'
            i += 1
            continue

        # H3
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.runs[0].font.name = 'Calibri'
            i += 1
            continue

        # Separador ---
        if stripped == '---':
            i += 1
            continue

        # Tabla
        if stripped.startswith('|'):
            rows, i = collect_table(lines, i)
            if not rows:
                continue
            n_cols = max(len(r) for r in rows)
            if n_cols <= 2:
                build_info_table(doc, rows)
            else:
                build_events_table(doc, rows)
            continue

        # Línea vacía
        if not stripped:
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        add_rich_text(p, stripped, font_size=10)
        i += 1


# ─── Estilos globales del documento ─────────────────────────────────────────
def configure_styles(doc):
    from docx.shared import Pt
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = hex_to_rgb(C_ROJO)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(13)
    h2.font.bold  = True
    h2.font.color.rgb = hex_to_rgb(C_OSCURO)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(11)
    h3.font.bold  = True
    h3.font.color.rgb = hex_to_rgb(C_ROJO)


# ─── Main ────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    configure_styles(doc)

    # Portada / título del documento
    cover = doc.add_heading('NARRATIVAS DE CASOS DE USO', level=0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('MOPGIMED — Sistema Inteligente de Inventario Farmacéutico')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = hex_to_rgb(C_OSCURO)
    doc.add_paragraph()

    files = sorted(glob.glob('NARRATIVAS-CU/narrativa_cu*.md'))
    print(f'Encontrados {len(files)} archivos de narrativa.')

    for idx, filepath in enumerate(files):
        fname = os.path.basename(filepath)
        print(f'  [{idx+1:02d}/25] {fname}')
        process_file(doc, filepath, is_first=(idx == 0))

    output = 'NARRATIVAS-MOPGIMED.docx'
    doc.save(output)
    print(f'\nDocumento generado: {output}')


if __name__ == '__main__':
    main()
