"""
Genera Sprint_Backlog_MOPGIMED.docx — versión consolidada (11 UCs + 7 RNFs).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colores ─────────────────────────────────────────────────────────────────
C_AZUL_HDR  = '1E3A5F'
C_AZUL_RF   = 'DBEAFE'
C_VERDE_RNF = 'DCFCE7'
C_SPRINT    = '1E40AF'
C_GRIS      = 'F1F5F9'
C_BLANCO    = 'FFFFFF'
C_ROJO      = 'B91C1C'
C_TEXTO     = '1E293B'

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_bg(cell, color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color)
    pr.append(shd)

def set_margins(cell, t=80, b=80, l=100, r=100):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for side, v in [('top', t), ('bottom', b), ('left', l), ('right', r)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(v)); el.set(qn('w:type'), 'dxa')
        m.append(el)
    pr.append(m)

def rgb(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def cell_write(cell, text, bold=False, size=9, color=None, center=False):
    set_margins(cell)
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size); run.font.name = 'Calibri'
    if color: run.font.color.rgb = rgb(color)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = 'Calibri'
    return p

def para(doc, text, size=10, bold=False, color=None, center=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(space_after)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color: r.font.color.rgb = rgb(color)
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = 'Calibri'


# ─── Datos ────────────────────────────────────────────────────────────────────

BACKLOG = [
    # ── Casos de Uso Funcionales ──────────────────────────────────────────────
    dict(id='UC-001', nombre='Iniciar sesión',
         desc='Permite validar credenciales y autorizar acceso según el rol asignado, garantizando que solo usuarios autorizados ingresen al sistema.',
         tipo='RF', prioridad='Alta',
         just='Es el punto de entrada al sistema; sin autenticación ninguna función está protegida. Todo el sistema depende de esta base.',
         owner='Salas Jiménez, W.', rfs='RF-001'),

    dict(id='UC-002', nombre='Gestionar usuarios y asignar rol',
         desc='Permite registrar, modificar, habilitar o deshabilitar usuarios y asignar perfiles de acceso (Administrador, Farmacia, Jefatura) directamente desde el formulario de creación o edición. Asignar rol es acción inherente a gestionar usuario.',
         tipo='RF', prioridad='Alta',
         just='Sin usuarios gestionados no existen roles ni permisos. RF-003 se integra en el mismo flujo de RF-002: el rol se asigna al crear o editar un usuario, no como acción separada.',
         owner='Vargas G., A. / Salas J., W.', rfs='RF-002, RF-003'),

    dict(id='UC-003', nombre='Gestión de medicamentos',
         desc='Permite administrar el ciclo completo del medicamento: listar el inventario, registrar nuevos medicamentos con metadatos de lote (costo, stock inicial, fecha de vencimiento), modificar registros existentes y consultar disponibilidad y precio de venta.',
         tipo='RF', prioridad='Alta',
         just='Módulo central del inventario. Agrupa las cuatro operaciones CRUD del medicamento. RF-008 (metadatos) es subformulario de RF-005; RF-009 (modificar) y RF-010 (consultar) operan sobre el mismo objeto.',
         owner='Vargas G., A. / Salas J., W.', rfs='RF-004, RF-005, RF-008, RF-009, RF-010'),

    dict(id='UC-004', nombre='Búsqueda y filtrado de medicamentos',
         desc='Permite buscar medicamentos mediante inteligencia artificial con sugerencias automáticas estandarizadas de fuentes farmacéuticas, y refinar los resultados aplicando filtros por laboratorio fabricante.',
         tipo='RF', prioridad='Alta',
         just='La búsqueda con IA (RF-006) y el filtro por laboratorio (RF-007) actúan sobre el mismo listado y comparten el mismo flujo. Filtrar por laboratorio es un modo especializado de búsqueda, no un módulo independiente.',
         owner='Vargas Gutierrez, A.', rfs='RF-006, RF-007'),

    dict(id='UC-005', nombre='Control de inventario',
         desc='Permite visualizar medicamentos próximos a vencer o ya vencidos, mostrar el estado del stock mediante semáforo visual (normal, bajo, crítico) y aplicar filtros por laboratorio, mes o año sobre la lista de vencimientos.',
         tipo='RF', prioridad='Alta',
         just='RF-013 (filtrar vencimientos) es subpaso directo de RF-011. RF-012 (semáforo) complementa la misma vista de inventario. Los tres forman un único módulo de control de stock y fechas.',
         owner='Salas Jiménez, W.', rfs='RF-011, RF-012, RF-013'),

    dict(id='UC-006', nombre='Gestión de notificaciones',
         desc='Permite generar notificaciones automáticas sobre bajo stock, stock crítico o vencimientos próximos, y visualizarlas en una bandeja consultable ordenada por prioridad.',
         tipo='RF', prioridad='Alta',
         just='Las alertas generadas (RF-014) son el contenido de la bandeja (RF-015). No existe bandeja funcional sin alertas. Separar "generar alerta" de "mostrarla en bandeja" crea una fragmentación artificial.',
         owner='Vargas Gutierrez, A.', rfs='RF-014, RF-015'),

    dict(id='UC-007', nombre='Reporte general de medicamentos',
         desc='Permite emitir reportes consolidados del inventario con filtros por periodo, costo de compra o venta, y exportar los resultados en formato Excel para análisis o presentación externa.',
         tipo='RF', prioridad='Alta',
         just='RF-017 (definir filtros) es un paso dentro del flujo de RF-016. RF-023 (exportar Excel) es la acción final del mismo reporte. Forman un único flujo: generar → filtrar → exportar.',
         owner='Vargas Gutierrez, A.', rfs='RF-016, RF-017, RF-023'),

    dict(id='UC-008', nombre='Reporte de ingresos y salidas',
         desc='Permite emitir reportes históricos sobre ingresos al almacén y salidas por ventas o consumo, con filtrado por rangos de tiempo definidos por el usuario.',
         tipo='RF', prioridad='Alta',
         just='RF-019 (filtrar por tiempo) es el mecanismo exclusivo de personalización del reporte de RF-018 y no aplica de forma independiente a otros módulos del sistema.',
         owner='Vargas Gutierrez, A.', rfs='RF-018, RF-019'),

    dict(id='UC-009', nombre='Reporte de rotación de medicamentos',
         desc='Permite consultar los medicamentos con mayor nivel de salida o venta dentro de un periodo determinado y generar el reporte correspondiente para su exportación y análisis.',
         tipo='RF', prioridad='Media',
         just='RF-021 (generar reporte) es la acción de formalizar y exportar lo que visualiza RF-020. Son la vista y la exportación del mismo análisis; separarlos no aporta valor funcional adicional.',
         owner='Salas Jiménez, W.', rfs='RF-020, RF-021'),

    dict(id='UC-010', nombre='Dashboard analítico',
         desc='Permite mostrar indicadores generales del inventario mediante gráficos y métricas de apoyo para la toma de decisiones gerenciales, incluyendo análisis avanzados con Machine Learning sobre demanda futura, medicamentos críticos y patrones estacionales de consumo.',
         tipo='RF', prioridad='Alta',
         just='El dashboard inteligente (RF-024) es la extensión con IA del dashboard base (RF-022). Son el mismo módulo con dos niveles de análisis: operativo e inteligente con ML. No tiene sentido presentarlos como módulos distintos.',
         owner='Vargas Gutierrez, A.', rfs='RF-022, RF-024'),

    dict(id='UC-011', nombre='Visualizar historial de acciones',
         desc='Permite consultar el registro de actividades realizadas por los usuarios del sistema para control de auditabilidad y seguimiento operativo interno.',
         tipo='RF', prioridad='Media',
         just='Apoya la auditabilidad interna pero no afecta la operación diaria del inventario. Se integra al módulo de administración como complemento de gestión de usuarios.',
         owner='Salas Jiménez, W.', rfs='RF-025'),

    # ── Requerimientos No Funcionales ─────────────────────────────────────────
    dict(id='RNF-001', nombre='Seguridad',
         desc='El sistema deberá garantizar acceso controlado mediante autenticación y roles, protegiendo la información de medicamentos y usuarios.',
         tipo='RNF', prioridad='Alta',
         just='Sin seguridad la información sensible queda expuesta. Es un requisito base para cualquier sistema de salud. Prioridad 1.',
         owner='Salas Jiménez, W.', rfs='RNF-001'),

    dict(id='RNF-002', nombre='Disponibilidad',
         desc='El sistema deberá estar disponible de forma continua para la consulta de medicamentos, alertas y reportes.',
         tipo='RNF', prioridad='Alta',
         just='El inventario necesita acceso permanente. Una caída puede afectar la atención al paciente. Prioridad 1.',
         owner='Vargas Gutierrez, A.', rfs='RNF-002'),

    dict(id='RNF-003', nombre='Consistencia de datos',
         desc='El sistema deberá mantener coherencia entre registros de inventario y su visualización en consultas, reportes y dashboard.',
         tipo='RNF', prioridad='Alta',
         just='Datos inconsistentes generan decisiones erróneas. Es fundamental para la confiabilidad del sistema. Prioridad 1.',
         owner='Salas Jiménez, W.', rfs='RNF-003'),

    dict(id='RNF-004', nombre='Rendimiento',
         desc='Las operaciones de búsqueda, filtrado y generación de reportes deberán ejecutarse en máximo 3 segundos en condiciones normales.',
         tipo='RNF', prioridad='Alta',
         just='La lentitud afecta directamente la productividad del personal y la continuidad operativa. Prioridad 1.',
         owner='Vargas Gutierrez, A.', rfs='RNF-004'),

    dict(id='RNF-005', nombre='Actualización de información',
         desc='Los cambios en el inventario deberán reflejarse de manera inmediata en todo el sistema, incluyendo alertas y paneles.',
         tipo='RNF', prioridad='Alta',
         just='La desactualización puede provocar errores de stock y decisiones basadas en datos obsoletos. Prioridad 1.',
         owner='Salas Jiménez, W.', rfs='RNF-005'),

    dict(id='RNF-006', nombre='Usabilidad',
         desc='La interfaz deberá permitir tareas como registro, consulta y filtrado de medicamentos de forma clara y sin complejidad.',
         tipo='RNF', prioridad='Media',
         just='Facilita la adopción del sistema pero no impide el funcionamiento técnico. Los módulos operan aunque la UI sea mejorable. Prioridad 2.',
         owner='Vargas Gutierrez, A.', rfs='RNF-006'),

    dict(id='RNF-007', nombre='Procesamiento analítico',
         desc='El sistema deberá procesar datos históricos para reportes y análisis del dashboard sin afectar el desempeño general.',
         tipo='RNF', prioridad='Media',
         just='Mejora el análisis estratégico pero el sistema base es funcional sin el componente predictivo activo. Prioridad 2.',
         owner='Salas Jiménez, W.', rfs='RNF-007'),
]

SPRINTS = [
    dict(
        num=1, fecha='11/05/2026', entregable=None,
        items=[
            dict(id='UC-001', nombre='Iniciar sesión',     owner='Salas Jiménez, W.',        tipo='RF'),
            dict(id='RNF-001', nombre='Seguridad',          owner='Salas Jiménez, W.',        tipo='RNF'),
        ],
        desc=(
            'El Sprint 1 constituye la base fundamental del sistema MOPGIMED. Se implementa UC-001 '
            '(Iniciar sesión), que valida credenciales contra la base de datos y retorna el menú '
            'personalizado según el rol del usuario (Administrador, Farmacia o Jefatura). De forma '
            'complementaria, se trabaja el requerimiento no funcional de Seguridad (RNF-001), '
            'asegurando control de acceso diferenciado por perfil. Sin estos dos elementos no es '
            'posible avanzar de forma segura hacia los módulos funcionales posteriores. La '
            'trazabilidad se sustentará con el SRS del módulo de autenticación, el SAD con la vista '
            'lógica y diagrama de componentes de Acceso al Sistema, la implementación del módulo '
            'Auth en backend y frontend, y los casos de prueba de login correcto e incorrecto.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=2, fecha='18/05/2026', entregable=None,
        items=[
            dict(id='UC-002', nombre='Gestionar usuarios y asignar rol', owner='Vargas G., A. / Salas J., W.', tipo='RF'),
            dict(id='UC-011', nombre='Visualizar historial de acciones',  owner='Salas Jiménez, W.',            tipo='RF'),
        ],
        desc=(
            'El Sprint 2 implementa la administración de usuarios del sistema. UC-002 (Gestionar '
            'usuarios y asignar rol) permite registrar nuevos usuarios, modificar sus datos, '
            'habilitar o deshabilitar cuentas, y asignar el rol correspondiente directamente desde '
            'el formulario de creación o edición, cubriendo en un único flujo los anteriores RF-002 '
            'y RF-003. UC-011 (Visualizar historial de acciones) registra y muestra el historial de '
            'actividades de cada usuario para control de auditabilidad, cubriendo el anterior RF-025. '
            'Ambos UCs pertenecen al módulo de administración y comparten la pantalla de Gestión de '
            'Usuarios. La trazabilidad incluirá el SRS del módulo de usuarios, el SAD con los '
            'diagramas de secuencia correspondientes, la implementación de los endpoints POST y PUT '
            'de usuarios, y pruebas de creación, edición y desactivación de cuentas.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=3, fecha='25/05/2026', entregable=None,
        items=[
            dict(id='UC-003',  nombre='Gestión de medicamentos', owner='Vargas G., A. / Salas J., W.', tipo='RF'),
            dict(id='RNF-002', nombre='Disponibilidad',           owner='Vargas Gutierrez, A.',         tipo='RNF'),
            dict(id='RNF-003', nombre='Consistencia de datos',    owner='Salas Jiménez, W.',            tipo='RNF'),
        ],
        desc=(
            'El Sprint 3 desarrolla el núcleo operativo del inventario farmacéutico mediante UC-003 '
            '(Gestión de medicamentos), que implementa las cuatro operaciones CRUD: listar el '
            'inventario completo, registrar nuevos medicamentos con todos sus metadatos (código, '
            'laboratorio, lote, stock, costo, precio de venta, fecha de vencimiento y registro '
            'sanitario), modificar registros existentes y consultar disponibilidad y precio de venta. '
            'Este UC consolida los anteriores RF-004, RF-005, RF-008, RF-009 y RF-010. RNF-002 '
            '(Disponibilidad) garantiza operación continua del sistema, y RNF-003 (Consistencia de '
            'datos) asegura que cualquier cambio se refleje coherentemente en todas las vistas. '
            'La trazabilidad comprenderá el SRS del módulo de inventario, el SAD con diagramas de '
            'secuencia de gestión de medicamentos, la implementación de los servicios CRUD y sus '
            'pruebas de creación, actualización y consulta de registros.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=4, fecha='01/06/2026', entregable='Entregable Unidad 2',
        items=[
            dict(id='UC-004',  nombre='Búsqueda y filtrado de medicamentos', owner='Vargas Gutierrez, A.', tipo='RF'),
            dict(id='UC-005',  nombre='Control de inventario',                owner='Salas Jiménez, W.',   tipo='RF'),
            dict(id='RNF-004', nombre='Rendimiento',                          owner='Vargas Gutierrez, A.', tipo='RNF'),
        ],
        desc=(
            'El Sprint 4 es el Entregable de Unidad 2. UC-004 (Búsqueda y filtrado) potencia la '
            'localización de medicamentos mediante búsqueda inteligente con IA que sugiere nombres '
            'estandarizados de fuentes farmacéuticas, complementada con filtro por laboratorio '
            'fabricante (consolida RF-006 y RF-007). UC-005 (Control de inventario) implementa el '
            'módulo de monitoreo de stock: visualización de medicamentos próximos a vencer o '
            'vencidos, semáforo de estado (normal/bajo/crítico) y filtros por laboratorio, mes o '
            'año sobre la lista de vencimientos (consolida RF-011, RF-012 y RF-013). RNF-004 '
            '(Rendimiento) asegura que todas las operaciones de búsqueda y filtrado respondan en '
            'menos de 3 segundos. La trazabilidad abarcará el SRS del motor de búsqueda IA y del '
            'módulo de inventario, el SAD con los diagramas de secuencia correspondientes, la '
            'implementación del componente de autocompletado y del módulo de vencimientos, y '
            'pruebas de tiempo de respuesta y precisión de sugerencias.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=5, fecha='08/06/2026', entregable='Entregable Unidad 3',
        items=[
            dict(id='UC-006',  nombre='Gestión de notificaciones',    owner='Vargas Gutierrez, A.', tipo='RF'),
            dict(id='RNF-005', nombre='Actualización de información', owner='Salas Jiménez, W.',   tipo='RNF'),
        ],
        desc=(
            'El Sprint 5 implementa el sistema de alertas proactivo del inventario mediante UC-006 '
            '(Gestión de notificaciones), que consolida los anteriores RF-014 y RF-015. Se desarrolla '
            'la generación automática de notificaciones ante eventos críticos (stock bajo, stock '
            'crítico, vencimientos próximos) y la bandeja de notificaciones donde el personal puede '
            'revisar todas las alertas activas ordenadas por prioridad, marcando las leídas de forma '
            'individual o masiva. RNF-005 (Actualización de información) garantiza que cualquier '
            'cambio en el inventario se propague inmediatamente a la bandeja y a los paneles del '
            'sistema. La trazabilidad incluirá el SRS del módulo de alertas, el SAD con el diagrama '
            'de secuencia de notificaciones, la implementación del generador de alertas, y pruebas '
            'de activación ante umbrales configurados de stock y vencimiento.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=6, fecha='15/06/2026', entregable='Entregable Unidad 3',
        items=[
            dict(id='UC-007',  nombre='Reporte general de medicamentos', owner='Vargas Gutierrez, A.', tipo='RF'),
            dict(id='RNF-006', nombre='Usabilidad',                       owner='Vargas Gutierrez, A.', tipo='RNF'),
        ],
        desc=(
            'El Sprint 6 construye el módulo de Reportes Operativos. UC-007 (Reporte general) '
            'implementa el flujo completo del reporte consolidado de inventario: generación del '
            'reporte con datos de todos los medicamentos, configuración de filtros por periodo o '
            'costos (consolida RF-016 y RF-017), y exportación en formato Excel para análisis '
            'externo (RF-023). El reporte, el filtrado y la exportación se integran en un único '
            'flujo coherente. RNF-006 (Usabilidad) guía el diseño de la interfaz de reportes para '
            'que sea clara e intuitiva, reduciendo la curva de aprendizaje del personal de farmacia '
            'y jefatura. La trazabilidad comprenderá el SRS del módulo de reportes, el SAD con el '
            'diagrama de componentes del motor de reportes, la implementación del servicio de '
            'generación y los filtros dinámicos, y pruebas con diferentes combinaciones de filtros.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=7, fecha='22/06/2026', entregable='Entregable Unidad 3',
        items=[
            dict(id='UC-008', nombre='Reporte de ingresos y salidas',        owner='Vargas Gutierrez, A.', tipo='RF'),
            dict(id='UC-009', nombre='Reporte de rotación de medicamentos',  owner='Salas Jiménez, W.',   tipo='RF'),
        ],
        desc=(
            'El Sprint 7 desarrolla las capacidades analíticas de movimiento del inventario. '
            'UC-008 (Reporte de ingresos y salidas) genera reportes históricos de entradas al '
            'almacén y salidas por ventas o consumo, con filtrado por rangos de tiempo definidos '
            'por el usuario (consolida RF-018 y RF-019). UC-009 (Reporte de rotación) visualiza '
            'los medicamentos con mayor nivel de salida en un periodo determinado y genera el '
            'reporte correspondiente para exportación y análisis (consolida RF-020 y RF-021). '
            'Ambos UCs transforman datos operativos en información estratégica para la toma de '
            'decisiones de compras y abastecimiento. La trazabilidad incluirá el SRS del módulo '
            'de análisis, el SAD con los diagramas de secuencia correspondientes, la implementación '
            'de los servicios de reporte y cálculo de rotación, y pruebas de precisión de datos.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
    dict(
        num=8, fecha='29/06/2026', entregable='Entregable Unidad 3 — Entregable Final',
        items=[
            dict(id='UC-010',  nombre='Dashboard analítico',      owner='Vargas Gutierrez, A.', tipo='RF'),
            dict(id='RNF-007', nombre='Procesamiento analítico',  owner='Salas Jiménez, W.',   tipo='RNF'),
        ],
        desc=(
            'El Sprint 8 es el Entregable Final del proyecto MOPGIMED. UC-010 (Dashboard analítico) '
            'implementa el módulo central de inteligencia del sistema en dos niveles: el dashboard '
            'base con indicadores generales del inventario (KPIs, gráficos de movimientos semanales, '
            'comparativa mensual, tendencias por categoría) y el dashboard inteligente con el modelo '
            'de Machine Learning para predicción de demanda futura, identificación de medicamentos en '
            'riesgo de desabasto y análisis de patrones estacionales de consumo (consolida RF-022 y '
            'RF-024). RNF-007 (Procesamiento analítico) garantiza que el modelo ML procese los datos '
            'históricos sin degradar el rendimiento general del sistema. Este sprint representa la '
            'propuesta de valor diferencial de MOPGIMED frente a sistemas de inventario tradicionales. '
            'La trazabilidad final abarcará el SRS completo del sistema, el SAD integral con todos '
            'los diagramas, la implementación del modelo predictivo y sus componentes de visualización, '
            'pruebas de precisión del modelo y rendimiento bajo carga, y evidencia de la funcionalidad '
            'completa del sistema integrado.'
        ),
        trazabilidad=[
            'Documento SRS: Sección de requerimientos cubiertos por el sprint.',
            'Documento SAD: Vistas arquitectónicas relacionadas (secuencia, componentes, clases).',
            'Implementación: Inventario de componentes creados que responden al SRS y SAD.',
            'Documento Evidencia Pruebas: Casos de prueba ejecutados con resultados.',
            'Evidencia de Funcionalidad en Ejecución: [Insertar captura de pantalla aquí]',
        ],
    ),
]

# ─── Generación del documento ─────────────────────────────────────────────────
def build_product_backlog(doc):
    heading(doc, 'III. Product Backlog – Inventario de Requerimientos', level=1)
    para(doc,
        'El Product Backlog contiene la totalidad de los requerimientos del sistema MOPGIMED, '
        'consolidados en 11 Casos de Uso funcionales y 7 Requerimientos No Funcionales (18 ítems '
        'en total). Los casos de uso funcionales se destacan en color azul y los requerimientos no '
        'funcionales en color verde. Cada ítem indica los RFs originales que agrupa, la prioridad '
        'asignada con su justificación y el Task Owner responsable.')

    cols = ['ID', 'Nombre', 'Descripción / RFs agrupados', 'Tipo', 'Prioridad', 'Justificación de Prioridad', 'Task Owner']
    widths = [1.4, 3.2, 5.5, 1.1, 1.3, 4.8, 2.8]

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    hrow = tbl.rows[0]
    for c_idx, col in enumerate(cols):
        set_bg(hrow.cells[c_idx], C_AZUL_HDR)
        cell_write(hrow.cells[c_idx], col, bold=True, size=9, color=C_BLANCO, center=True)

    # Datos
    for item in BACKLOG:
        row = tbl.add_row()
        bg = C_AZUL_RF if item['tipo'] == 'RF' else C_VERDE_RNF
        desc_full = f"{item['desc']}\n→ Agrupa: {item['rfs']}" if item['tipo'] == 'RF' else item['desc']
        values = [item['id'], item['nombre'], desc_full, item['tipo'],
                  item['prioridad'], item['just'], item['owner']]
        for c_idx, val in enumerate(values):
            set_bg(row.cells[c_idx], bg)
            cell_write(row.cells[c_idx], val, size=8,
                       center=(c_idx in [0, 3, 4]))

    # Anchos
    for row in tbl.rows:
        for c_idx, w in enumerate(widths):
            row.cells[c_idx].width = Cm(w)

    doc.add_paragraph()
    para(doc, '■  Azul = Caso de Uso Funcional (UC)          ■  Verde = Requerimiento No Funcional (RNF)',
         size=9, bold=True)
    doc.add_paragraph()


def build_sprint(doc, sprint, idx):
    titulo = f"{'IV' if idx == 0 else str(idx + 4)}."  # lettering simple
    num   = sprint['num']
    fecha = sprint['fecha']
    ent   = sprint['entregable']

    label = f"Sprint Backlog  [Fecha: {fecha}]"
    if ent:
        label += f"  ⟶  {ent}"

    heading(doc, label, level=2)

    # Tabla del sprint
    cols = ['ID', 'Nombre', 'Task Owner', 'Tipo']
    widths = [1.8, 6.5, 4.5, 1.5]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hrow = tbl.rows[0]
    for c_idx, col in enumerate(cols):
        set_bg(hrow.cells[c_idx], C_SPRINT)
        cell_write(hrow.cells[c_idx], col, bold=True, size=9, color=C_BLANCO, center=True)

    for item in sprint['items']:
        row = tbl.add_row()
        bg = C_AZUL_RF if item['tipo'] == 'RF' else C_VERDE_RNF
        for c_idx, val in enumerate([item['id'], item['nombre'], item['owner'], item['tipo']]):
            set_bg(row.cells[c_idx], bg)
            cell_write(row.cells[c_idx], val, size=9, center=(c_idx in [0, 3]))

    for row in tbl.rows:
        for c_idx, w in enumerate(widths):
            row.cells[c_idx].width = Cm(w)

    doc.add_paragraph()

    # Descripción
    para(doc, 'Descripción del Sprint:', bold=True, size=10)
    para(doc, sprint['desc'], size=10)

    # Trazabilidad
    para(doc, 'Trazabilidad requerida:', bold=True, size=10)
    for item_t in sprint['trazabilidad']:
        bullet(doc, item_t, size=10)

    doc.add_paragraph()


def configure_styles(doc):
    styles = doc.styles
    n = styles['Normal']
    n.font.name = 'Calibri'; n.font.size = Pt(10)

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'; h1.font.size = Pt(14)
    h1.font.bold = True; h1.font.color.rgb = rgb(C_AZUL_HDR)

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'; h2.font.size = Pt(12)
    h2.font.bold = True; h2.font.color.rgb = rgb(C_SPRINT)


def main():
    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    configure_styles(doc)

    # ── Portada ───────────────────────────────────────────────────────────────
    for line in [
        'UNIVERSIDAD PRIVADA DE TACNA',
        'FACULTAD DE INGENIERÍA',
        'Escuela Profesional de Ingeniería de Sistemas',
    ]:
        p = para(doc, line, size=12, bold=True, center=True, space_after=6)

    doc.add_paragraph()
    para(doc,
         '"Modelo predictivo basado en Machine Learning para la gestión de inventario '
         'de medicamentos en la Clínica La Luz, 2026"',
         size=13, bold=True, center=True, space_after=12)

    doc.add_paragraph()
    for lbl, val in [
        ('Curso:', 'Construcción de Software I'),
        ('Docente:', 'Mag. Ricardo Eduardo Valcárcel Alvarado'),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(lbl + '  '); r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
        r2 = p.add_run(val);       r2.font.size = Pt(11); r2.font.name = 'Calibri'

    doc.add_paragraph()
    para(doc, 'Integrantes:', size=11, bold=True, center=True)
    para(doc, 'Salas Jiménez, Walter Emmanuel          (2022073896)', size=11, center=True)
    para(doc, 'Vargas Gutierrez, Angel Jose              (2020066922)', size=11, center=True)
    doc.add_paragraph()
    para(doc, 'Tacna – Perú', size=11, center=True)
    para(doc, '2026 - I', size=11, center=True)

    doc.add_page_break()

    # ── Control de versiones ──────────────────────────────────────────────────
    heading(doc, 'CONTROL DE VERSIONES', level=1)
    cv = doc.add_table(rows=2, cols=6)
    cv.style = 'Table Grid'
    for c_idx, h in enumerate(['Versión','Hecha por','Revisada por','Aprobada por','Fecha','Motivo']):
        set_bg(cv.rows[0].cells[c_idx], C_AZUL_HDR)
        cell_write(cv.rows[0].cells[c_idx], h, bold=True, size=9, color=C_BLANCO, center=True)
    vals = ['1.0','AJVG, WESJ','-','-','06/05/2026','Versión 1.0']
    for c_idx, v in enumerate(vals):
        cell_write(cv.rows[1].cells[c_idx], v, size=9, center=True)
    doc.add_paragraph()

    doc.add_page_break()

    # ── Título principal ──────────────────────────────────────────────────────
    para(doc, 'Modelo predictivo basado en Machine Learning\npara la gestión de inventario de medicamentos\nen la Clínica La Luz, 2026\n(MOPGIMED)',
         size=14, bold=True, center=True, space_after=10)
    para(doc, 'Planificación de Sprints Backlog', size=12, bold=True, center=True)
    para(doc, 'Versión 1.0', size=11, center=True)

    doc.add_page_break()

    # ── Introducción ──────────────────────────────────────────────────────────
    heading(doc, 'I. Introducción', level=1)
    para(doc,
        'El presente Documento de Planificación del Proyecto corresponde al sistema MOPGIMED '
        '(Modelo Predictivo de Gestión de Inventario de Medicamentos), desarrollado para la '
        'Clínica La Luz como parte del curso Construcción de Software I. Este documento establece '
        'el marco metodológico de desarrollo ágil adoptado por el equipo, describiendo de manera '
        'estructurada la totalidad de los requerimientos identificados durante la etapa de análisis, '
        'consolidados en 11 Casos de Uso funcionales y 7 Requerimientos No Funcionales (18 ítems en '
        'el Product Backlog), así como su distribución planificada a lo largo de ocho sprints de trabajo.')
    para(doc,
        'El proyecto surge de la necesidad de optimizar los procesos de gestión de inventario de '
        'medicamentos dentro de la clínica, los cuales actualmente se realizan de forma manual, '
        'generando inconsistencias en los registros, dificultades en el control de stock, duplicidad '
        'de información y limitada capacidad de respuesta ante vencimientos o desabastecimientos. '
        'Para dar solución a esta problemática, se propone una plataforma web inteligente que integra '
        'módulos de registro y control de inventario, monitoreo mediante alertas automáticas, '
        'generación de reportes consolidados y un componente diferencial de inteligencia artificial '
        'basado en Machine Learning para la predicción de demanda y patrones de consumo.')
    para(doc,
        'La planificación sigue el marco de trabajo Scrum, organizando el trabajo en sprints '
        'semanales de duración fija, con entregables definidos para cada semana. Los 25 requerimientos '
        'funcionales originales fueron consolidados en 11 Casos de Uso agrupando aquellos que '
        'pertenecen al mismo módulo o cuya relación es de subpaso o extensión, eliminando fragmentación '
        'innecesaria sin perder cobertura funcional. El equipo de desarrollo está compuesto por '
        'Salas Jiménez, Walter Emmanuel y Vargas Gutierrez, Angel Jose, quienes distribuyen la carga '
        'de trabajo de manera equitativa entre los módulos del sistema.')

    doc.add_page_break()

    # ── Objetivo ──────────────────────────────────────────────────────────────
    heading(doc, 'II. Objetivo', level=1)
    heading(doc, '1. Objetivo General', level=2)
    para(doc,
        'Establecer una hoja de ruta clara, organizada y trazable para el desarrollo del sistema '
        'MOPGIMED, que permita al equipo ejecutar el proyecto de manera ordenada, con compromisos '
        'semanales verificables y con los entregables de cada sprint debidamente sustentados '
        'mediante la documentación técnica y evidencias de funcionalidad.')

    heading(doc, '2. Objetivos Específicos', level=2)
    for obj in [
        'Inventariar de manera completa y categorizada los requerimientos del sistema, consolidando '
        'los 25 RFs originales en 11 Casos de Uso funcionales, diferenciando los funcionales de los '
        'no funcionales, asignando prioridades justificadas y responsables de implementación para cada uno.',

        'Definir la distribución de los 18 ítems del backlog en ocho sprints semanales, garantizando '
        'una carga de trabajo equilibrada de 2 a 3 ítems por sprint, seleccionados de forma lógica '
        'según su nivel de dependencia, criticidad operativa y complejidad técnica.',

        'Describir con detalle suficiente el alcance de cada sprint, especificando qué se implementará, '
        'por qué se priorizó de esa manera y cuál será la evidencia de cumplimiento, incluyendo la '
        'trazabilidad con el SRS, el SAD, la implementación, las pruebas y la evidencia de '
        'funcionalidad en ejecución.',

        'Garantizar la alineación entre la planificación inicial y los artefactos técnicos producidos '
        'durante cada iteración, favoreciendo la calidad del producto final y la trazabilidad integral '
        'del proceso de construcción del software.',
    ]:
        bullet(doc, obj)

    doc.add_page_break()

    # ── Product Backlog ───────────────────────────────────────────────────────
    build_product_backlog(doc)

    doc.add_page_break()

    # ── Sprint Backlog ────────────────────────────────────────────────────────
    heading(doc, 'IV. Sprint Backlog – Requerimientos por Sprint', level=1)
    para(doc,
        'A continuación se presenta el Sprint Backlog del proyecto MOPGIMED. Se han definido 8 '
        'sprints semanales con 2 a 3 ítems por sprint (18 ítems en total). Para cada sprint se '
        'detalla la fecha de presentación, los casos de uso comprometidos y una descripción extendida '
        'del alcance, criterios de entrega y trazabilidad requerida.')
    doc.add_paragraph()

    for idx, sprint in enumerate(SPRINTS):
        build_sprint(doc, sprint, idx)
        if idx < len(SPRINTS) - 1:
            doc.add_paragraph()

    output = 'Sprint_Backlog_MOPGIMED.docx'
    doc.save(output)
    print(f'Documento generado: {output}')


if __name__ == '__main__':
    main()
