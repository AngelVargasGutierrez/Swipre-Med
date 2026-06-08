"""
Genera INVENTARIO-IMPLEMENTACION-MOPGIMED.docx — Punto 3 de trazabilidad del sprint.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generar_srs_epis import (
    hex_to_rgb, C_ROJO, C_OSCURO,
    _build_data_table, _add_heading_paragraph, _add_body_paragraph,
)


def set_margins(doc, cm_val=2.5):
    for section in doc.sections:
        section.top_margin = Cm(cm_val)
        section.bottom_margin = Cm(cm_val)
        section.left_margin = Cm(cm_val)
        section.right_margin = Cm(cm_val)
from trazabilidad_mopgimed import INVENTARIO_IMPLEMENTACION


def build():
    doc = Document()
    set_margins(doc, 2.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('INVENTARIO DE COMPONENTES — MOPGIMED')
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = hex_to_rgb(C_ROJO)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Trazabilidad: Implementación ↔ SRS ↔ SAD')
    r2.font.size = Pt(11)
    r2.font.color.rgb = hex_to_rgb(C_OSCURO)
    doc.add_paragraph()

    _add_body_paragraph(doc, (
        'Documento de soporte al Sprint Backlog. Lista los componentes implementados en el '
        'repositorio MOPGIMED (frontend, backend y base de datos) y su relación con los '
        'requerimientos funcionales del SRS y las vistas del SAD.'
    ))

    rows = [(c, t, r, rf, uc, s) for c, t, r, rf, uc, s in INVENTARIO_IMPLEMENTACION]
    _build_data_table(
        doc,
        ['Componente', 'Tipo', 'Ruta / Archivo', 'RF', 'UC', 'Referencia SAD'],
        rows,
        [2.2, 1.3, 3.8, 1.5, 1.5, 3.7],
    )

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'INVENTARIO-IMPLEMENTACION-MOPGIMED.docx')
    doc.save(out)
    print(f'[OK] {out} ({len(rows)} componentes)')


if __name__ == '__main__':
    build()
