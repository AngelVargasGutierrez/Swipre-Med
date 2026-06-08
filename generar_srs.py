"""
generar_srs.py
Genera FD03-SRS-MOPGIMED.docx — Especificación de Requerimientos de Software
Sistema Inteligente de Inventario Farmacéutico — MOPGIMED
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
#  COLORES
# ─────────────────────────────────────────────────────────────────────────────
C_ROJO         = 'B91C1C'   # H1
C_OSCURO       = '1E293B'   # H2
C_AZUL_OSCURO  = '1E3A5F'   # encabezados de tabla
C_AZUL_CLARO   = 'DBEAFE'   # filas RF
C_VERDE_CLARO  = 'DCFCE7'   # filas RNF
C_GRIS_CLARO   = 'F1F5F9'   # campo label
C_GRIS_CODE    = 'F1F5F9'   # bloques de código
C_BLANCO       = 'FFFFFF'
C_NEGRO        = '111827'


# ─────────────────────────────────────────────────────────────────────────────
#  UTILIDADES
# ─────────────────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_bg(cell, color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color.lstrip('#'))
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def cell_write(cell, text, bold=False, size=10, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p   = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = hex_to_rgb(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size      = Pt(14)
        run.font.color.rgb = hex_to_rgb(C_ROJO)
    elif level == 2:
        run.font.size      = Pt(12)
        run.font.color.rgb = hex_to_rgb(C_OSCURO)
    else:
        run.font.size      = Pt(11)
        run.font.color.rgb = hex_to_rgb(C_OSCURO)
    return p


def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = hex_to_rgb(C_NEGRO)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold       = True
    run.italic     = True
    run.font.name  = 'Calibri'
    run.font.size  = Pt(9)
    run.font.color.rgb = hex_to_rgb(C_OSCURO)


def add_code_block(doc, code_text):
    """Inserta el código PlantUML en una tabla con fondo gris."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_bg(cell, C_GRIS_CODE)
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    '100')
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7)
    run.font.color.rgb = hex_to_rgb('1E293B')
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def set_margins(doc, cm_val=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(cm_val)
        section.bottom_margin = Cm(cm_val)
        section.left_margin   = Cm(cm_val)
        section.right_margin  = Cm(cm_val)


# ─────────────────────────────────────────────────────────────────────────────
#  DATOS DEL PROYECTO
# ─────────────────────────────────────────────────────────────────────────────
TEAM_MEMBERS = [
    "Vargas Gutierrez, Angel",
    "Lopez Mamani, Carlos",
    "Quispe Flores, Maria",
    "Condori Tapia, Luis",
]

CURSO    = "Ingeniería de Software I — SI885"
DOCENTE  = "Mg. Ing. Rosa María Flores Condori"
YEAR     = "2025"
UNIV     = "Universidad Privada de Tacna"
FACULTAD = "Facultad de Ingeniería"
ESCUELA  = "Escuela Profesional de Ingeniería de Sistemas"

# ─────────────────────────────────────────────────────────────────────────────
#  CONTENIDO — PLANTUML INLINES
# ─────────────────────────────────────────────────────────────────────────────

PUML_PAQUETES_SRS = """\
@startuml MOPGIMED Diagrama de Paquetes SRS
top to bottom direction
title "MOPGIMED — Diagrama de Paquetes\\nEspecificación de Requerimientos (SRS)"
skinparam {
  backgroundColor        #FFFFFF
  shadowing              false
  defaultFontName        Segoe UI
  PackageBackgroundColor #FFF5F5
  PackageBorderColor     #B91C1C
  PackageBorderThickness 2
  PackageFontSize        11
  PackageFontStyle       Bold
  PackageFontColor       #7F1D1D
  ComponentBorderColor   #FECACA
  ArrowColor             #B91C1C
  ArrowThickness         1.5
}
package "Acceso al Sistema" as P1 {
  [UC-001: Iniciar Sesion]
}
package "Administracion de Usuarios" as P2 {
  [UC-002: Gestionar Usuarios y Asignar Rol]
  [UC-011: Visualizar Historial de Acciones]
}
package "Gestion de Medicamentos" as P3 {
  [UC-003: Gestion de Medicamentos - CRUD]
  [UC-004: Busqueda y Filtrado de Medicamentos]
}
package "Control de Inventario" as P4 {
  [UC-005: Control de Inventario]
}
package "Alertas y Notificaciones" as P5 {
  [UC-006: Gestion de Notificaciones]
}
package "Reportes" as P6 {
  [UC-007: Reporte General de Medicamentos]
  [UC-008: Reporte de Ingresos y Salidas]
  [UC-009: Reporte de Rotacion de Medicamentos]
}
package "Dashboard Analitico" as P7 {
  [UC-010: Dashboard Analitico]
}
P1 ..> P2 : autoriza al Administrador
P1 ..> P3 : autoriza a Farmacia y Admin
P1 ..> P4 : autoriza segun rol
P1 ..> P5 : autoriza a todos los roles
P1 ..> P6 : autoriza a Jefatura y Admin
P1 ..> P7 : autoriza a Jefatura y Admin
P3 ..> P4 : stock alimenta el inventario
P3 ..> P5 : cambios de stock generan alertas
P4 ..> P5 : alertas de stock bajo
P3 ..> P6 : datos para reportes
P4 ..> P6 : movimientos para reportes
P4 ..> P7 : indicadores para dashboard
P6 ..> P7 : datos agregados para analisis
@enduml"""

PUML_CASOS_USO = """\
@startuml DiagramaCasosDeUso MOPGIMED
title "Diagrama de Casos de Uso — MOPGIMED\\nSistema Inteligente de Inventario Farmacéutico"
left to right direction
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam ArrowColor #B91C1C
skinparam ArrowThickness 1.2
skinparam actor {
  BorderColor    #B91C1C
  BackgroundColor #FFF1F2
  FontColor      #7F1D1D
  FontStyle      Bold
}
skinparam usecase {
  BorderColor    #475569
  BackgroundColor #FFFFFF
  FontColor      #1E293B
}
actor "Administrador" as Admin
actor "Farmacia" as Farm
actor "Jefatura" as Jef
package "Acceso al Sistema" #EFF6FF {
  usecase "UC-001: Iniciar Sesion" as UC01
}
package "Administracion de Usuarios" #FFFBEB {
  usecase "UC-002: Gestionar Usuarios\\ny Asignar Rol" as UC02
  usecase "UC-011: Visualizar Historial\\nde Acciones" as UC11
}
package "Gestion de Medicamentos" #ECFDF5 {
  usecase "UC-003: Gestion de\\nMedicamentos - CRUD" as UC03
  usecase "UC-004: Busqueda y Filtrado\\nde Medicamentos" as UC04
}
package "Control de Inventario" #FFF7ED {
  usecase "UC-005: Control de Inventario" as UC05
}
package "Alertas y Notificaciones" #FFF1F2 {
  usecase "UC-006: Gestion de Notificaciones" as UC06
}
package "Reportes" #F5F3FF {
  usecase "UC-007: Reporte General\\nde Medicamentos" as UC07
  usecase "UC-008: Reporte de Ingresos\\ny Salidas" as UC08
  usecase "UC-009: Reporte de Rotacion\\nde Medicamentos" as UC09
}
package "Dashboard Analitico" #F0FDF4 {
  usecase "UC-010: Dashboard Analitico" as UC10
}
Admin --> UC01
Admin --> UC02
Admin --> UC11
Admin --> UC03
Admin --> UC04
Admin --> UC05
Admin --> UC06
Admin --> UC07
Admin --> UC08
Admin --> UC09
Admin --> UC10
Farm --> UC01
Farm --> UC03
Farm --> UC04
Jef --> UC01
Jef --> UC05
Jef --> UC06
Jef --> UC07
Jef --> UC08
Jef --> UC09
Jef --> UC10
@enduml"""

PUML_ROBUSTEZ_CU001 = """\
@startuml Robustez CU001 MOPGIMED
top to bottom direction
title "Diagrama de Robustez — CU001: Iniciar Sesion\\nMOPGIMED · Sistema Inteligente de Inventario Farmacéutico"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ArrowColor #B91C1C  ArrowThickness 1.5
}
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam boundary { BorderColor #1D4ED8  BackgroundColor #EFF6FF  FontColor #1E3A8A }
skinparam control { BorderColor #B45309  BackgroundColor #FFFBEB  FontColor #78350F }
skinparam entity { BorderColor #065F46  BackgroundColor #ECFDF5  FontColor #064E3B }
skinparam package { BorderColor #CBD5E1  BackgroundColor #F8FAFC  FontStyle Bold }
actor "Personal\\ndel Hospital" as ACTOR
package "Entidad" as PKG_E {
  entity "Registro del Personal\\nUsuario · Contrasena · Cargo" as E1
  entity "Permisos por Cargo\\nAdmin · Farmacia · Jefatura" as E2
  entity "Historial de Actividad\\nFecha · Hora · Usuario" as E3
  E1 -[hidden]right-> E2
  E2 -[hidden]right-> E3
}
package "Frontera" as PKG_B {
  boundary "Formulario de Inicio de Sesion\\nCampos: usuario y contrasena" as B1
  boundary "Aviso de Acceso Denegado\\nError: usuario o clave invalida" as B3
  boundary "Panel Principal personalizado\\nMenu segun cargo asignado" as B2
  B1 -[hidden]down-> B3
  B3 -[hidden]down-> B2
}
package "Control" as PKG_C {
  control "Verificador de Identidad\\nRevisa credenciales y estado Activo" as C1
  control "Cargador de Opciones por Cargo\\nPrepara menu y secciones" as C2
  control "Registrador de Actividad\\nAnota quien ingreso y cuando" as C3
  C1 -[hidden]down-> C2
  C2 -[hidden]down-> C3
}
PKG_E -[hidden]down-> PKG_B
PKG_B -[hidden]right-> PKG_C
ACTOR --> B1 : 1. Ingresa usuario y contrasena
B1   --> C1  : 2. Envia datos para verificacion
C1   --> E1  : 3. Busca usuario en el registro
C1   --> B3  : 4a. Datos incorrectos o inactivo
B3   --> ACTOR : muestra aviso de error
C1   --> C2  : 4b. Acceso valido
C2   --> E2  : 5. Consulta permisos del cargo
C1   --> C3  : 6. Notifica el ingreso
C3   --> E3  : 7. Guarda fecha, hora y usuario
C2   --> B2  : 8. Abre el panel personalizado
B2   --> ACTOR : 9. El personal navega el sistema
@enduml"""

PUML_ROBUSTEZ_CU002 = """\
@startuml Robustez CU002 MOPGIMED
top to bottom direction
title "Diagrama de Robustez — CU002: Gestionar Usuarios y Asignar Rol\\nMOPGIMED · Sistema Inteligente de Inventario Farmacéutico"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ArrowColor #B91C1C  ArrowThickness 1.5
}
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam boundary { BorderColor #1D4ED8  BackgroundColor #EFF6FF  FontColor #1E3A8A }
skinparam control { BorderColor #B45309  BackgroundColor #FFFBEB  FontColor #78350F }
skinparam entity { BorderColor #065F46  BackgroundColor #ECFDF5  FontColor #064E3B }
skinparam package { BorderColor #CBD5E1  BackgroundColor #F8FAFC  FontStyle Bold }
actor "Administrador\\ndel Sistema" as ACTOR
package "Entidad" as PKG_E {
  entity "Registro de Usuarios\\nNombre · Username · Email · Rol · Estado" as E1
  entity "Tabla de Roles\\nAdmin · Farmacia · Jefatura" as E2
  entity "Historial de Actividad\\nFecha · Administrador · Accion realizada" as E3
  E1 -[hidden]right-> E2
  E2 -[hidden]right-> E3
}
package "Frontera" as PKG_B {
  boundary "Pantalla Gestion de Usuarios\\nTabla con nombre, rol, estado y acciones" as B1
  boundary "Formulario Nuevo / Editar Usuario\\nCampos: nombre, username, contrasena, rol, email" as B2
  boundary "Aviso de Error\\nDatos incompletos o username duplicado" as B3
  B1 -[hidden]down-> B2
  B2 -[hidden]down-> B3
}
package "Control" as PKG_C {
  control "Gestor de Usuarios\\nCarga lista, crea, actualiza y cambia estado" as C1
  control "Validador de Datos\\nVerifica campos, username y contrasenas" as C2
  control "Asignador de Rol\\nDetermina etiqueta y permisos del rol" as C3
  control "Registrador de Actividad\\nGuarda quien, que operacion y cuando" as C4
  C1 -[hidden]down-> C2
  C2 -[hidden]down-> C3
  C3 -[hidden]down-> C4
}
PKG_E -[hidden]down-> PKG_B
PKG_B -[hidden]right-> PKG_C
ACTOR --> B1  : 1. Accede al modulo de Usuarios
B1    --> C1  : 2. Solicita la lista de usuarios
C1    --> E1  : 3. Consulta todos los usuarios
E1    --> B1  : 4. Muestra tabla con nombre y rol
ACTOR --> B2  : 5. Abre formulario nuevo o editar
B2    --> C2  : 6. Envia datos para validacion
C2    --> B3  : 7a. Datos invalidos o username duplicado
B3    --> ACTOR : muestra aviso de error
C2    --> C3  : 7b. Datos correctos, asigna rol
C3    --> E2  : 8. Consulta etiqueta del rol
C3    --> C1  : 9. Rol asignado, procede a guardar
C1    --> E1  : 10. Inserta o actualiza el registro
C1    --> C4  : 11. Notifica la accion realizada
C4    --> E3  : 12. Guarda en el historial
C1    --> B1  : 13. Actualiza la tabla de usuarios
B1    --> ACTOR : 14. Verifica el resultado
@enduml"""

PUML_ROBUSTEZ_CU011 = """\
@startuml Robustez CU011 MOPGIMED
top to bottom direction
title "Diagrama de Robustez — CU011: Visualizar Historial de Acciones\\nMOPGIMED · Sistema Inteligente de Inventario Farmacéutico"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ArrowColor #B91C1C  ArrowThickness 1.5
}
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam boundary { BorderColor #1D4ED8  BackgroundColor #EFF6FF  FontColor #1E3A8A }
skinparam control { BorderColor #B45309  BackgroundColor #FFFBEB  FontColor #78350F }
skinparam entity { BorderColor #065F46  BackgroundColor #ECFDF5  FontColor #064E3B }
skinparam package { BorderColor #CBD5E1  BackgroundColor #F8FAFC  FontStyle Bold }
actor "Administrador\\ndel Sistema" as ACTOR
package "Entidad" as PKG_E {
  entity "Historial de Actividad\\nFecha · Hora · Usuario · Accion · Modulo · Detalle" as E1
  entity "Registro de Usuarios\\nNombre · Rol (referencia para historial)" as E2
  E1 -[hidden]right-> E2
}
package "Frontera" as PKG_B {
  boundary "Pantalla Gestion de Usuarios\\nModulo principal con tabla de usuarios" as B1
  boundary "Tabla Historial de Acciones\\nFecha/Hora · Usuario · Accion · Modulo · Detalle" as B2
  B1 -[hidden]down-> B2
}
package "Control" as PKG_C {
  control "Consultor de Historial\\nRecupera registros ordenados por fecha desc" as C1
  control "Registrador de Actividad\\nInserta registro al acceder o gestionar usuarios" as C2
  C1 -[hidden]down-> C2
}
PKG_E -[hidden]down-> PKG_B
PKG_B -[hidden]right-> PKG_C
ACTOR --> B1  : 1. Accede al modulo Gestion de Usuarios
B1    --> C1  : 2. Solicita el historial de acciones
C1    --> E1  : 3. Consulta registros por fecha desc
E1    --> C1  : 4. Devuelve lista con acciones registradas
C1    --> B2  : 5. Retorna los registros de historial
B2    --> ACTOR : 6. Muestra tabla con todas las acciones
ACTOR --> B1  : 7. Realiza accion (login o gestion)
B1    --> C2  : 8. Sistema confirma la accion
C2    --> E2  : 9. Consulta nombre y rol del usuario
C2    --> E1  : 10. Inserta registro automaticamente
E1    --> C2  : 11. Confirma la insercion
C2    --> B1  : 12. Accion queda trazada en el historial
@enduml"""

PUML_SECUENCIA_CU001 = """\
@startuml Secuencia CU001 MOPGIMED
title "Diagrama de Secuencia — CU001: Iniciar Sesion\\nMOPGIMED · Sistema Inteligente de Inventario Farmacéutico"
autonumber
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam sequenceArrowColor #B91C1C
skinparam sequenceArrowThickness 1.5
skinparam sequenceLifeLineBorderColor #B91C1C
skinparam sequenceGroupBorderColor #B91C1C
skinparam sequenceGroupBackgroundColor #FFF5F5
actor "Personal del Hospital" as ACTOR
box "Frontera" #EFF6FF
  participant "Pantalla de\\nInicio de Sesion" as B1
  participant "Aviso de\\nAcceso Denegado" as B3
  participant "Panel Principal\\npersonalizado" as B2
end box
box "Control" #FFFBEB
  participant "Verificador\\nde Identidad" as C1
  participant "Cargador de Opciones\\npor Cargo" as C2
  participant "Registrador\\nde Actividad" as C3
end box
box "Entidad" #ECFDF5
  participant "Registro\\ndel Personal" as E1
  participant "Permisos\\npor Cargo" as E2
  participant "Historial\\nde Actividad" as E3
end box
ACTOR -> B1 : Abre el sistema e ingresa usuario y contrasena
B1 -> C1 : Envia los datos para verificacion
C1 -> E1 : Busca al usuario en el registro del personal
E1 --> C1 : Devuelve nombre, cargo y estado del usuario
alt Usuario o contrasena incorrectos / usuario Inactivo
  C1 -> B3 : Los datos no son validos
  B3 -> ACTOR : Muestra el aviso de error
else El usuario existe, la clave es correcta y esta Activo
  C1 -> C2 : Identidad confirmada, carga opciones del cargo
  C2 -> E2 : Consulta que puede ver segun su cargo
  E2 --> C2 : Devuelve el menu correspondiente al cargo
  C1 -> C3 : Notifica que el personal acaba de ingresar
  C3 -> E3 : Guarda fecha, hora y nombre del usuario
  C2 -> B2 : Abre el sistema con el menu personalizado
  B2 -> ACTOR : El personal ve su panel y usa el sistema
end
@enduml"""

PUML_SECUENCIA_CU002 = """\
@startuml Secuencia CU002 MOPGIMED
title "Diagrama de Secuencia — CU002: Gestionar Usuarios y Asignar Rol\\nMOPGIMED · Sistema Inteligente de Inventario Farmacéutico"
autonumber
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam sequenceArrowColor #B91C1C
skinparam sequenceArrowThickness 1.5
skinparam sequenceLifeLineBorderColor #B91C1C
skinparam sequenceGroupBorderColor #B91C1C
skinparam sequenceGroupBackgroundColor #FFF5F5
actor "Administrador\\ndel Sistema" as ACTOR
box "Frontera" #EFF6FF
  participant "Pantalla de\\nGestion de Usuarios" as B1
  participant "Formulario\\nNuevo / Editar Usuario" as B2
  participant "Aviso de\\nError de Datos" as B3
end box
box "Control" #FFFBEB
  participant "Gestor\\nde Usuarios" as C1
  participant "Validador\\nde Datos" as C2
  participant "Asignador\\nde Rol" as C3
  participant "Registrador\\nde Actividad" as C4
end box
box "Entidad" #ECFDF5
  participant "Registro\\nde Usuarios" as E1
  participant "Tabla\\nde Roles" as E2
  participant "Historial\\nde Actividad" as E3
end box
== Consultar lista de usuarios ==
ACTOR -> B1 : Accede al modulo Gestion de Usuarios
B1 -> C1 : Solicita la lista de usuarios registrados
C1 -> E1 : Consulta todos los usuarios del sistema
E1 --> C1 : Devuelve nombre, rol, estado y fecha creacion
C1 --> B1 : Retorna lista de usuarios
B1 --> ACTOR : Muestra tabla con todos los usuarios
== Crear nuevo usuario con rol ==
ACTOR -> B1 : Hace clic en "Nuevo Usuario"
B1 -> B2 : Abre formulario vacio de registro
B2 --> ACTOR : Muestra campos: nombre, username, contrasena, rol, email
ACTOR -> B2 : Completa campos y selecciona el rol
B2 -> C2 : Envia datos para validacion previa
alt Datos incompletos o username ya existe
  C2 -> E1 : Verifica si el username ya esta registrado
  E1 --> C2 : Confirma username duplicado o faltan campos
  C2 -> B3 : Datos invalidos
  B3 --> ACTOR : Muestra mensaje de error
else Datos validos y username disponible
  C2 -> C3 : Datos correctos, procede a asignar el rol
  C3 -> E2 : Consulta la etiqueta del rol seleccionado
  E2 --> C3 : Devuelve la etiqueta del rol
  C3 -> C1 : Rol asignado, procede a crear el usuario
  C1 -> E1 : Inserta nuevo usuario con rol y estado Activo
  E1 --> C1 : Confirma insercion con ID generado
  C1 -> C4 : Notifica que se creo un nuevo usuario
  C4 -> E3 : Guarda: quien creo, nombre, fecha y hora
  C1 --> B1 : Actualiza la lista con el nuevo usuario
  B1 --> ACTOR : Muestra el usuario creado en la tabla
end
@enduml"""

PUML_SECUENCIA_CU011 = """\
@startuml Secuencia CU011 MOPGIMED
title "Diagrama de Secuencia — CU011: Visualizar Historial de Acciones\\nMOPGIMED · Sistema Inteligente de Inventario Farmacéutico"
autonumber
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam sequenceArrowColor #B91C1C
skinparam sequenceArrowThickness 1.5
skinparam sequenceLifeLineBorderColor #B91C1C
skinparam sequenceGroupBorderColor #B91C1C
skinparam sequenceGroupBackgroundColor #FFF5F5
actor "Administrador\\ndel Sistema" as ACTOR
box "Frontera" #EFF6FF
  participant "Pantalla de\\nGestion de Usuarios" as B1
  participant "Tabla de\\nHistorial de Acciones" as B2
end box
box "Control" #FFFBEB
  participant "Consultor\\nde Historial" as C1
  participant "Registrador\\nde Actividad" as C2
end box
box "Entidad" #ECFDF5
  participant "Historial\\nde Actividad" as E1
  participant "Registro\\nde Usuarios" as E2
end box
== Visualizar historial de acciones ==
ACTOR -> B1 : Accede al modulo Gestion de Usuarios
B1 -> C1 : Solicita el historial de acciones del sistema
C1 -> E1 : Consulta los ultimos registros por fecha desc
E1 --> C1 : Devuelve lista: fecha, usuario, accion, modulo, detalle
C1 --> B2 : Retorna los registros de historial
B2 --> ACTOR : Muestra tabla con todas las acciones registradas
alt No hay registros en el historial
  C1 --> B2 : Lista vacia
  B2 --> ACTOR : Muestra: aun no hay acciones registradas
else Historial con registros disponibles
  B2 --> ACTOR : El administrador consulta el detalle de cada accion
end
== Registro automatico al gestionar usuarios ==
ACTOR -> B1 : Realiza una accion sobre un usuario
B1 -> C2 : Sistema procesa la accion solicitada
C2 -> E2 : Consulta nombre y rol del usuario
E2 --> C2 : Devuelve nombre y etiqueta de rol
C2 -> E1 : Inserta registro: admin, accion, modulo, usuario afectado, fecha
E1 --> C2 : Confirma insercion en el historial
C2 --> B1 : Registro guardado, la accion queda trazada
@enduml"""


# ─────────────────────────────────────────────────────────────────────────────
#  REQUERIMIENTOS FUNCIONALES (17 RF)
#  Tuplas: (id, requerimiento, pantalla_actor, prioridad, task_owner)
# ─────────────────────────────────────────────────────────────────────────────
RF_SECTIONS = [
    ("Seguridad y administración", [
        ("RF-001", "Iniciar sesión",
         "Login — Todos los roles", "Alta", "Salas Jiménez, W."),
        ("RF-002", "Gestionar usuarios y asignar rol",
         "Usuarios (+ modal) — Administrador", "Alta", "Vargas G., A. / Salas J., W."),
        ("RF-003", "Visualizar historial de acciones",
         "Historial (sección en Usuarios) — Administrador", "Media", "Salas Jiménez, W."),
    ]),
    ("Gestión de medicamentos — Farmacia / Almacén", [
        ("RF-004", "Registrar lote de medicamento existente (lote, fecha venc., unidad, cant. por envase)",
         "NuevoMedicamento — Farmacia / Almacén", "Alta", "Vargas Gutierrez, A."),
        ("RF-005", "Registrar medicamento nuevo con asistencia IA (búsqueda sugerida, selección y metadatos de lote)",
         "NuevoMedicamento — Farmacia / Almacén", "Alta", "Vargas Gutierrez, A."),
        ("RF-006", "Consultar medicamento (disponibilidad y precio)",
         "Consulta — Caja farmacia", "Alta", "Vargas Gutierrez, A."),
    ]),
    ("Gestión de medicamentos — Jefatura / Logística", [
        ("RF-007", "Listar medicamentos (buscador + exportar Excel)",
         "Medicamentos — Jefatura / Logística", "Alta", "Salas Jiménez, W."),
        ("RF-008", "Filtrar listado de medicamentos (por laboratorio u otros criterios)",
         "Medicamentos — Jefatura / Logística", "Alta", "Salas Jiménez, W."),
        ("RF-009", "Modificar medicamento",
         "Medicamentos — Jefatura / Logística", "Alta", "Vargas Gutierrez, A."),
    ]),
    ("Control de inventario", [
        ("RF-010", "Visualizar semáforo de stock de medicamentos",
         "ControlInventario — Jefatura / Logística", "Alta", "Salas Jiménez, W."),
        ("RF-011", "Visualizar semáforo de medicamentos por vencer / vencidos",
         "ControlInventario — Jefatura / Logística", "Alta", "Salas Jiménez, W."),
        ("RF-012", "Listar medicamentos con semáforo (estado normal, bajo, crítico)",
         "ControlInventario — Jefatura / Logística", "Alta", "Salas Jiménez, W."),
        ("RF-013", "Filtrar lista de vencimientos (por laboratorio, mes o año)",
         "ControlInventario — Jefatura / Logística", "Media", "Salas Jiménez, W."),
    ]),
    ("Reportes", [
        ("RF-014", "Reporte general de medicamentos (filtros + exportar Excel)",
         "Reportes → general — Jefatura", "Alta", "Vargas Gutierrez, A."),
        ("RF-015", "Reporte de ingresos y salidas (filtro por tiempo)",
         "Reportes → movimientos — Jefatura", "Alta", "Vargas Gutierrez, A."),
        ("RF-016", "Reporte de rotación de medicamentos",
         "Reportes → rotación — Jefatura", "Media", "Salas Jiménez, W."),
    ]),
    ("Dashboard analítico", [
        ("RF-017", "Visualizar dashboard analítico (operativo + inteligente ML/IA)",
         "Dashboard / Analytics — Jefatura / Logística", "Alta", "Vargas Gutierrez, A."),
    ]),
]

# ─────────────────────────────────────────────────────────────────────────────
#  DATOS DE CASOS DE USO
# ─────────────────────────────────────────────────────────────────────────────
UCS = [
    {
        "id": "UC-001", "nombre": "Iniciar Sesión", "rf": "RF-001",
        "actores": "Personal del Hospital (Administrador, Jefatura, Farmacia)",
        "descripcion": (
            "El sistema valida las credenciales ingresadas y otorga acceso a los módulos según "
            "el cargo del usuario. Administrador: acceso completo. Jefatura: reportes, análisis "
            "e indicadores. Farmacia: inventario, medicamentos y alertas."
        ),
        "precondicion": "El usuario debe estar registrado y con estado 'Activo' en la base de datos.",
        "postcondicion": (
            "Sesión iniciada correctamente. El sistema redirige al Panel Principal según el cargo "
            "y registra el ingreso en el Historial de Actividad."
        ),
        "flujo_principal": [
            "El personal abre el sistema MOPGIMED en su navegador.",
            "El sistema muestra la Pantalla de Inicio de Sesión con los campos Usuario y Contraseña.",
            "El personal ingresa su nombre de usuario y contraseña.",
            "El personal presiona el botón 'Iniciar Sesión'.",
            "El sistema envía las credenciales al Verificador de Identidad.",
            "El Verificador busca al usuario en el Registro del Personal.",
            "El sistema comprueba que la contraseña coincida con la almacenada.",
            "El sistema verifica que el estado del usuario sea 'Activo'.",
            "El sistema identifica el cargo y carga los módulos correspondientes.",
            "El Registrador de Actividad guarda en el Historial: usuario, cargo, fecha y hora.",
            "El sistema muestra el Panel Principal con el menú personalizado.",
        ],
        "flujos_alternos": [
            ("Credenciales incorrectas", "El sistema muestra el aviso: 'Usuario o contraseña incorrectos.' El personal corrige e intenta de nuevo."),
            ("Usuario Inactivo", "El sistema bloquea el acceso y muestra: 'Su cuenta está inactiva. Comuníquese con el Administrador.'"),
        ],
        "ui_detail": True,
    },
    {
        "id": "UC-002", "nombre": "Gestionar Usuarios y Asignar Rol", "rf": "RF-002",
        "actores": "Administrador",
        "descripcion": (
            "El sistema permite al Administrador gestionar de forma completa las cuentas de acceso. "
            "Puede registrar nuevos usuarios, consultar la lista, editar información y cambiar el estado "
            "de las cuentas. Toda acción queda registrada en el Historial de Actividad."
        ),
        "precondicion": "El Administrador debe haber iniciado sesión con rol Administrador.",
        "postcondicion": (
            "El usuario queda registrado, actualizado, habilitado o deshabilitado en la base de datos. "
            "La acción queda registrada en el Historial de Actividad."
        ),
        "flujo_principal": [
            "El Administrador hace clic en 'Usuarios' del menú lateral izquierdo.",
            "El sistema carga el módulo y consulta todos los usuarios registrados.",
            "El sistema muestra la tabla de usuarios con: nombre, username, rol, estado y fecha.",
            "El Administrador hace clic en 'Nuevo Usuario' para registrar una cuenta nueva.",
            "El sistema abre el Formulario de Usuario con los campos vacíos.",
            "El Administrador completa: nombre, username, contraseña, rol y email.",
            "El Administrador selecciona el rol: Administrador, Jefatura o Farmacia.",
            "El Administrador hace clic en 'Guardar'.",
            "El sistema valida los datos: campos obligatorios, username único, contraseñas coincidentes.",
            "El sistema guarda el nuevo usuario en la base de datos con estado 'Activo'.",
            "El Registrador guarda en el Historial: Administrador, acción, usuario, fecha y hora.",
            "El sistema cierra el formulario y muestra el usuario creado en la tabla.",
        ],
        "flujos_alternos": [
            ("Datos incompletos o username duplicado", "El sistema no guarda ningún dato. Muestra el mensaje de error y resalta con borde rojo los campos con problema. El Administrador corrige y vuelve a guardar."),
            ("Editar usuario existente", "El Administrador hace clic en el ícono lápiz. El formulario se abre con los datos pre-cargados. Al guardar, el sistema aplica las mismas validaciones y actualiza el registro."),
            ("Cambiar estado del usuario", "El Administrador hace clic en el ícono de estado. El sistema cambia el estado a Activo o Inactivo según corresponda y registra la acción en el historial."),
        ],
        "ui_detail": True,
    },
    {
        "id": "UC-003", "nombre": "Gestión de Medicamentos - CRUD", "rf": "RF-004, RF-005, RF-007, RF-009",
        "actores": "Administrador, Farmacia",
        "descripcion": (
            "El sistema proporciona un módulo central para administrar la información de medicamentos. "
            "El personal puede registrar, consultar, editar y eliminar fichas que contienen: código, nombre, "
            "laboratorio, lote, stock, precio, vencimiento, categoría y estado calculado automáticamente."
        ),
        "precondicion": "El personal debe haber iniciado sesión con rol Administrador o Farmacia.",
        "postcondicion": "El medicamento queda registrado, actualizado o eliminado en la base de datos con su estado de stock calculado.",
        "flujo_principal": [
            "El personal hace clic en 'Medicamentos' del menú lateral.",
            "El sistema carga y muestra la lista completa de medicamentos.",
            "El sistema calcula automáticamente el estado de stock de cada medicamento.",
            "El personal hace clic en 'Nuevo Medicamento' para registrar una ficha nueva.",
            "El sistema muestra el formulario con todos los campos requeridos.",
            "El personal completa los campos: código, nombre, laboratorio, lote, stock, precio, vencimiento.",
            "El personal hace clic en 'Guardar'.",
            "El sistema valida los datos y calcula el estado del stock.",
            "El sistema guarda la ficha en la base de datos.",
            "El Registrador guarda la acción en el Historial de Actividad.",
        ],
        "flujos_alternos": [
            ("Datos incompletos o inválidos", "El sistema muestra el mensaje de error y mantiene el formulario abierto para corrección."),
            ("Editar medicamento", "El personal hace clic en Editar. El formulario se pre-carga con los datos actuales. Al guardar, el sistema recalcula el estado del stock."),
            ("Eliminar medicamento", "El sistema solicita confirmación. Al confirmar, elimina el registro y registra la acción en el historial."),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-004", "nombre": "Búsqueda y Filtrado de Medicamentos", "rf": "RF-006, RF-008",
        "actores": "Administrador, Farmacia",
        "descripcion": (
            "El sistema permite al personal buscar medicamentos por nombre, laboratorio o categoría "
            "en tiempo real, con filtrado combinado y resultados instantáneos."
        ),
        "precondicion": "El personal debe haber iniciado sesión en MOPGIMED. Deben existir medicamentos registrados.",
        "postcondicion": "El sistema muestra la lista filtrada de medicamentos que coinciden con el criterio de búsqueda.",
        "flujo_principal": [
            "El personal accede al módulo Medicamentos.",
            "El sistema muestra la lista completa de medicamentos.",
            "El personal escribe en el campo de búsqueda.",
            "El sistema filtra en tiempo real por nombre, laboratorio y categoría.",
            "El personal visualiza los resultados coincidentes.",
        ],
        "flujos_alternos": [
            ("Sin resultados", "El sistema muestra: 'No se encontraron medicamentos con ese criterio.' El personal limpia el campo para restaurar la lista completa."),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-005", "nombre": "Control de Inventario", "rf": "RF-010, RF-011, RF-012, RF-013",
        "actores": "Administrador, Jefatura, Farmacia",
        "descripcion": (
            "El sistema proporciona un módulo de control de inventario con semáforo de stock que clasifica "
            "los medicamentos en Normal, Bajo o Crítico según el stock actual versus el stock mínimo. "
            "También muestra los medicamentos próximos a vencer."
        ),
        "precondicion": "El personal debe haber iniciado sesión. Deben existir medicamentos registrados con stock mínimo definido.",
        "postcondicion": "El personal visualiza el estado del inventario con el semáforo de stock y las alertas de vencimiento.",
        "flujo_principal": [
            "El personal hace clic en 'Control de Inventario' del menú lateral.",
            "El sistema carga todos los medicamentos con su estado de stock calculado.",
            "El sistema muestra el semáforo: verde (Normal), amarillo (Bajo), rojo (Crítico).",
            "El sistema muestra la lista de medicamentos próximos a vencer en los próximos 30 días.",
            "El personal puede filtrar por estado de stock o por fecha de vencimiento.",
        ],
        "flujos_alternos": [
            ("Sin medicamentos críticos", "El sistema muestra el mensaje: 'No hay medicamentos en estado crítico actualmente.'"),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-006", "nombre": "Gestión de Notificaciones", "rf": "—",
        "actores": "Administrador, Jefatura, Farmacia",
        "descripcion": (
            "El sistema genera notificaciones automáticas por stock bajo, stock crítico y medicamentos "
            "próximos a vencer. El personal puede ver la bandeja de notificaciones, marcar como leídas "
            "individualmente o todas a la vez."
        ),
        "precondicion": "El personal debe haber iniciado sesión. El sistema debe haber generado al menos una notificación.",
        "postcondicion": "Las notificaciones quedan marcadas como leídas en la base de datos.",
        "flujo_principal": [
            "El personal hace clic en 'Notificaciones' del menú lateral o en el ícono de campana.",
            "El sistema muestra la bandeja de notificaciones con todas las alertas clasificadas por prioridad.",
            "El sistema resalta con color rojo las notificaciones de alta prioridad y azul las de baja.",
            "El personal lee las notificaciones de interés.",
            "El personal hace clic en 'Marcar como leída' para una notificación específica.",
            "El sistema actualiza el estado de la notificación en la base de datos.",
            "El personal puede hacer clic en 'Marcar todas como leídas' para limpiar la bandeja.",
        ],
        "flujos_alternos": [
            ("Sin notificaciones pendientes", "El sistema muestra: 'No tienes notificaciones pendientes.'"),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-007", "nombre": "Reporte General de Medicamentos", "rf": "RF-014",
        "actores": "Administrador, Jefatura",
        "descripcion": (
            "El sistema genera un reporte general del inventario con indicadores de valor, cantidad por "
            "estado y distribución por categoría. Incluye opción de exportación a Excel."
        ),
        "precondicion": "El usuario debe tener rol Administrador o Jefatura.",
        "postcondicion": "El sistema muestra el reporte general. Opcionalmente se descarga el archivo Excel.",
        "flujo_principal": [
            "El usuario accede al módulo Reportes.",
            "El sistema carga y calcula los indicadores del inventario.",
            "El sistema muestra: total de productos, valor de inventario, ganancia estimada, productos por estado.",
            "El sistema muestra la distribución de medicamentos por categoría.",
            "El usuario puede hacer clic en 'Exportar a Excel' para descargar el reporte.",
        ],
        "flujos_alternos": [
            ("Sin datos en el inventario", "El sistema muestra los indicadores en cero e informa que no hay medicamentos registrados."),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-008", "nombre": "Reporte de Ingresos y Salidas", "rf": "RF-015",
        "actores": "Administrador, Jefatura",
        "descripcion": (
            "El sistema genera un reporte de movimientos del inventario mostrando los ingresos y salidas "
            "de medicamentos por semana y la comparativa mensual de ventas versus compras."
        ),
        "precondicion": "El usuario debe tener rol Administrador o Jefatura. Deben existir movimientos registrados.",
        "postcondicion": "El sistema muestra el reporte de ingresos y salidas con gráficas de movimientos.",
        "flujo_principal": [
            "El usuario accede al módulo Reportes y selecciona 'Ingresos y Salidas'.",
            "El sistema consulta los movimientos de la semana actual.",
            "El sistema muestra el gráfico de barras con entradas y salidas por día.",
            "El sistema muestra la comparativa mensual de ventas versus compras.",
            "El usuario puede filtrar por rango de fechas.",
        ],
        "flujos_alternos": [
            ("Sin movimientos registrados", "El sistema muestra el gráfico vacío e informa que no hay movimientos en el período seleccionado."),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-009", "nombre": "Reporte de Rotación de Medicamentos", "rf": "RF-016",
        "actores": "Administrador, Jefatura",
        "descripcion": (
            "El sistema genera un reporte de los medicamentos de mayor rotación, permitiendo identificar "
            "los más demandados y exportar el informe para análisis externo."
        ),
        "precondicion": "El usuario debe tener rol Administrador o Jefatura.",
        "postcondicion": "El sistema muestra el reporte de rotación. Opcionalmente se exporta a Excel.",
        "flujo_principal": [
            "El usuario accede al módulo Reportes y selecciona 'Rotación de Medicamentos'.",
            "El sistema consulta y clasifica los medicamentos por nivel de rotación.",
            "El sistema muestra la lista ordenada de los medicamentos de mayor a menor rotación.",
            "El usuario puede filtrar por período de tiempo.",
            "El usuario puede exportar el reporte a Excel.",
        ],
        "flujos_alternos": [
            ("Sin datos de rotación", "El sistema informa que no hay datos de rotación para el período seleccionado."),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-010", "nombre": "Dashboard Analítico", "rf": "RF-017",
        "actores": "Administrador, Jefatura",
        "descripcion": (
            "El sistema muestra un panel analítico con indicadores clave de inventario, predicción de "
            "demanda por IA, tendencias por categoría y productos en riesgo de desabasto."
        ),
        "precondicion": "El usuario debe tener rol Administrador o Jefatura.",
        "postcondicion": "El sistema muestra el dashboard actualizado con los últimos datos del inventario.",
        "flujo_principal": [
            "El usuario accede al módulo Dashboard.",
            "El sistema carga todos los indicadores: total medicamentos, valor inventario, alertas activas.",
            "El sistema muestra el gráfico de predicción de demanda por mes.",
            "El sistema muestra las tendencias de consumo por categoría.",
            "El sistema lista los productos en riesgo de desabasto con prioridad Alta, Media o Baja.",
            "El sistema muestra las alertas activas en el panel lateral.",
        ],
        "flujos_alternos": [
            ("Error de conexión", "El sistema muestra un mensaje de error y sugiere recargar la página."),
        ],
        "ui_detail": False,
    },
    {
        "id": "UC-011", "nombre": "Visualizar Historial de Acciones", "rf": "RF-003",
        "actores": "Administrador",
        "descripcion": (
            "El sistema registra automáticamente cada acción relevante (inicio de sesión, creación, "
            "edición o cambio de estado de usuarios) y permite al Administrador consultar el historial "
            "completo con fecha, usuario, acción, módulo y detalle."
        ),
        "precondicion": "El Administrador debe haber iniciado sesión con rol Administrador.",
        "postcondicion": "El Administrador visualiza el historial de acciones. No se realizan cambios en la base de datos.",
        "flujo_principal": [
            "El Administrador accede al módulo Gestión de Usuarios.",
            "El sistema muestra la tabla de usuarios en la parte superior.",
            "En la sección inferior, el sistema carga el historial de acciones.",
            "El sistema consulta los registros ordenados por fecha descendente.",
            "El sistema muestra la tabla con: Fecha/Hora, Usuario, Acción, Módulo (badge), Detalle.",
            "El Administrador revisa el historial para auditar las acciones realizadas.",
        ],
        "flujos_alternos": [
            ("Sin registros en el historial", "El sistema muestra: 'Aún no hay acciones registradas en el sistema.'"),
        ],
        "ui_detail": True,
    },
]

RNF_DATA = [
    ("RNF-001", "Seguridad", "El sistema debe autenticar a todos los usuarios mediante credenciales únicas (username y contraseña). El acceso a cada módulo está restringido según el rol asignado. Las contraseñas se almacenan de forma segura.", "Alta"),
    ("RNF-002", "Disponibilidad", "El sistema debe estar disponible durante el horario operativo del hospital (07:00–22:00 horas) con un tiempo de inactividad no mayor a 30 minutos por mantenimiento.", "Alta"),
    ("RNF-003", "Consistencia de datos", "Todas las operaciones de escritura (registro, edición, eliminación) deben garantizar la integridad referencial de la base de datos MySQL. No se permiten registros huérfanos.", "Alta"),
    ("RNF-004", "Rendimiento", "El sistema debe cargar cualquier módulo en menos de 3 segundos en condiciones normales de red local. Las consultas a la base de datos no deben superar 1 segundo de tiempo de respuesta.", "Media"),
    ("RNF-005", "Actualización", "El estado del stock (Normal, Bajo, Crítico) y las notificaciones deben calcularse y actualizarse automáticamente cada vez que se guarda o edita un medicamento.", "Alta"),
    ("RNF-006", "Usabilidad", "La interfaz debe ser intuitiva y no requerir capacitación técnica previa. Los colores del semáforo y los badges de estado deben ser comprensibles sin necesidad de leyenda adicional.", "Media"),
    ("RNF-007", "Procesamiento analítico", "El módulo Analytics debe mostrar predicciones de demanda, tendencias y alertas críticas basadas en los datos del inventario con un margen de error menor al 15%.", "Media"),
]

REGLAS_NEGOCIO = [
    ("RN-001", "Cálculo de estado de stock", "El estado del stock se calcula automáticamente: Normal si stock > stockMínimo; Bajo si stock <= stockMínimo y stock > 50% de stockMínimo; Crítico si stock <= 50% de stockMínimo."),
    ("RN-002", "Unicidad de username", "El nombre de usuario (username) debe ser único en el sistema. El sistema rechazará el registro o edición de un usuario si el username ya existe en la base de datos."),
    ("RN-003", "Control de acceso por rol", "Solo el Administrador puede acceder al módulo de Gestión de Usuarios y al Historial de Acciones. Los roles Farmacia y Jefatura no tienen acceso a estos módulos."),
    ("RN-004", "Registro automático de historial", "Toda acción de creación, edición, activación o desactivación de usuarios, así como el inicio de sesión de cualquier usuario, debe registrarse automáticamente en la tabla historial_acciones."),
    ("RN-005", "Bloqueo de usuario inactivo", "Un usuario con estado 'Inactivo' no puede iniciar sesión en el sistema aunque sus credenciales sean correctas. El sistema mostrará un mensaje indicando que la cuenta está deshabilitada."),
    ("RN-006", "Validación de vencimientos", "El sistema debe alertar sobre medicamentos cuya fecha de vencimiento se encuentre dentro de los próximos 30 días. Estos medicamentos aparecen destacados en el módulo de Control de Inventario."),
    ("RN-007", "Integridad de datos de medicamentos", "No se pueden registrar medicamentos con stock negativo, precio de venta cero ni fecha de vencimiento anterior a la fecha actual. El sistema validará estos campos antes de guardar."),
]


# ─────────────────────────────────────────────────────────────────────────────
#  CONSTRUCCIÓN DEL DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────
def build_srs():
    doc = Document()
    set_margins(doc, 2.5)

    # Fuente base global
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    def center(text, size=10, bold=False, color=C_NEGRO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.color.rgb = hex_to_rgb(color)
        return p

    center(UNIV,    14, True,  C_OSCURO)
    center(FACULTAD, 12, False, C_OSCURO)
    center(ESCUELA,  11, False, C_OSCURO)
    doc.add_paragraph()
    center("FD03 — Especificación de Requerimientos de Software", 16, True, C_ROJO)
    center("Sistema Inteligente de Inventario Farmacéutico", 14, True, C_OSCURO)
    center("MOPGIMED", 18, True, C_ROJO)
    doc.add_paragraph()
    center(f"Curso: {CURSO}",   11, False, C_OSCURO)
    center(f"Docente: {DOCENTE}", 11, False, C_OSCURO)
    doc.add_paragraph()
    center("Integrantes:", 11, True, C_OSCURO)
    for m in TEAM_MEMBERS:
        center(m, 10, False, C_OSCURO)
    doc.add_paragraph()
    center(f"Tacna, {YEAR}", 10, False, C_OSCURO)

    page_break(doc)

    # ── I. INTRODUCCIÓN ───────────────────────────────────────────────────────
    add_heading(doc, "I. INTRODUCCIÓN", 1)

    add_heading(doc, "1.1 Generalidades de la Empresa", 2)
    add_body(doc, (
        "La Clínica La Luz es una institución de salud privada ubicada en la ciudad de Tacna, Perú. "
        "Cuenta con un área de Farmacia que gestiona el inventario de medicamentos utilizados en los "
        "diferentes servicios clínicos. Actualmente el proceso de control de inventario se realiza de "
        "forma manual o con herramientas de propósito general, lo que genera ineficiencias en el registro, "
        "seguimiento y control del stock farmacéutico."
    ))

    add_heading(doc, "1.2 Nombre, Visión y Misión", 2)
    add_body(doc, "Nombre del sistema: MOPGIMED — Sistema Inteligente de Inventario Farmacéutico")
    add_body(doc, (
        "Visión: Ser el sistema de referencia para la gestión inteligente de inventarios farmacéuticos "
        "en clínicas privadas de la región de Tacna, integrando inteligencia artificial para la predicción "
        "de demanda y la prevención de desabasto de medicamentos esenciales."
    ))
    add_body(doc, (
        "Misión: Proporcionar a la Clínica La Luz una herramienta web segura, eficiente e intuitiva que "
        "permita al personal de farmacia, jefatura y administración gestionar el inventario farmacéutico "
        "en tiempo real, con alertas automáticas, reportes detallados y análisis predictivo basado en IA."
    ))

    add_heading(doc, "1.3 Descripción del Problema", 2)
    add_body(doc, (
        "El área de Farmacia de la Clínica La Luz enfrenta los siguientes problemas en la gestión del "
        "inventario farmacéutico:"
    ))
    problems = [
        "Falta de visibilidad en tiempo real del stock de medicamentos, lo que genera situaciones de desabasto no detectadas a tiempo.",
        "Ausencia de un sistema de alertas automáticas para medicamentos con stock crítico o próximos a vencer.",
        "Registros manuales propensos a errores humanos en el control de ingresos y salidas.",
        "Imposibilidad de generar reportes automatizados para la toma de decisiones por parte de la Jefatura.",
        "Carencia de herramientas de predicción de demanda para anticipar necesidades de reabastecimiento.",
    ]
    for prob in problems:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(prob)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

    add_heading(doc, "1.4 Objetivos del Sistema", 2)
    objectives = [
        "Implementar un sistema web de control de inventario farmacéutico en tiempo real con acceso por roles.",
        "Generar alertas automáticas de stock crítico, stock bajo y medicamentos próximos a vencer.",
        "Proveer reportes detallados de ingresos, salidas y rotación de medicamentos exportables a Excel.",
        "Integrar un módulo de análisis con inteligencia artificial para la predicción de demanda farmacéutica.",
        "Registrar automáticamente todas las acciones del personal en un historial de actividad auditable.",
    ]
    for obj in objectives:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(obj)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

    add_heading(doc, "1.5 Alcance", 2)
    add_body(doc, (
        "MOPGIMED es un sistema web desarrollado con React 19 + Vite (frontend) y Node.js + Express (backend), "
        "con base de datos MySQL. El sistema cubre los siguientes módulos funcionales: autenticación de usuarios "
        "con control de acceso por roles, gestión completa de medicamentos (CRUD), control de inventario con "
        "semáforo de stock, centro de notificaciones, generación de reportes con exportación a Excel, dashboard "
        "analítico con predicción de demanda por IA, y administración de usuarios con historial de acciones. "
        "El sistema está diseñado para operar en red local del hospital y es accesible desde cualquier navegador web moderno."
    ))

    page_break(doc)

    # ── II. ESPECIFICACIÓN DE REQUERIMIENTOS ──────────────────────────────────
    add_heading(doc, "II. ESPECIFICACIÓN DE REQUERIMIENTOS", 1)

    add_heading(doc, "2.1 Requerimientos No Funcionales (RNF)", 2)
    add_body(doc, "A continuación se presentan los 7 requerimientos no funcionales del sistema MOPGIMED:")

    rnf_table = doc.add_table(rows=1, cols=4)
    rnf_table.style = 'Table Grid'
    rnf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs = ["ID", "Nombre", "Descripción", "Prioridad"]
    widths = [Cm(1.8), Cm(3.5), Cm(9.5), Cm(2.2)]
    for i, (hdr, w) in enumerate(zip(hdrs, widths)):
        cell = rnf_table.rows[0].cells[i]
        cell.width = w
        set_bg(cell, C_AZUL_OSCURO)
        cell_write(cell, hdr, bold=True, size=9, color=C_BLANCO,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for rnf_id, rnf_name, rnf_desc, rnf_prio in RNF_DATA:
        row = rnf_table.add_row()
        for i, val in enumerate([rnf_id, rnf_name, rnf_desc, rnf_prio]):
            cell = row.cells[i]
            set_bg(cell, C_VERDE_CLARO)
            cell_write(cell, val, size=9)
    doc.add_paragraph()

    add_heading(doc, "2.2 Requerimientos Funcionales (RF)", 2)
    add_body(doc, (
        "El sistema MOPGIMED define 17 requerimientos funcionales, agrupados por módulo e interfaz. "
        "Cada RF corresponde a una pantalla o flujo de usuario identificable en la aplicación."
    ))

    rf_hdrs = ["ID", "Requerimiento", "Pantalla / Actor", "Prioridad", "Task Owner"]
    rf_widths = [Cm(1.6), Cm(5.8), Cm(4.2), Cm(1.8), Cm(3.6)]

    for section_title, section_rows in RF_SECTIONS:
        add_heading(doc, section_title, 3)
        rf_table = doc.add_table(rows=1, cols=5)
        rf_table.style = 'Table Grid'
        rf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (hdr, w) in enumerate(zip(rf_hdrs, rf_widths)):
            cell = rf_table.rows[0].cells[i]
            cell.width = w
            set_bg(cell, C_AZUL_OSCURO)
            cell_write(cell, hdr, bold=True, size=8, color=C_BLANCO,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for row_data in section_rows:
            row = rf_table.add_row()
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                set_bg(cell, C_AZUL_CLARO)
                cell_write(cell, val, size=8)
        doc.add_paragraph()

    add_heading(doc, "2.3 Reglas de Negocio", 2)
    rn_table = doc.add_table(rows=1, cols=3)
    rn_table.style = 'Table Grid'
    rn_hdrs = ["ID", "Nombre", "Descripción"]
    rn_widths = [Cm(1.8), Cm(4.5), Cm(10.7)]
    for i, (hdr, w) in enumerate(zip(rn_hdrs, rn_widths)):
        cell = rn_table.rows[0].cells[i]
        cell.width = w
        set_bg(cell, C_AZUL_OSCURO)
        cell_write(cell, hdr, bold=True, size=9, color=C_BLANCO,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, (rn_id, rn_name, rn_desc) in enumerate(REGLAS_NEGOCIO):
        row = rn_table.add_row()
        bg = C_GRIS_CLARO if idx % 2 == 0 else C_BLANCO
        for i, val in enumerate([rn_id, rn_name, rn_desc]):
            cell = row.cells[i]
            set_bg(cell, bg)
            cell_write(cell, val, size=9)
    doc.add_paragraph()

    page_break(doc)

    # ── III. FASE DE DESARROLLO ───────────────────────────────────────────────
    add_heading(doc, "III. FASE DE DESARROLLO", 1)
    add_heading(doc, "3.1 Perfiles de Usuario", 2)

    perfiles = [
        ("Administrador", [
            "Acceso completo a todos los módulos del sistema.",
            "Módulo exclusivo: Gestión de Usuarios (crear, editar, activar/desactivar cuentas).",
            "Módulo exclusivo: Historial de Acciones (auditoría de todas las acciones del sistema).",
            "Acceso a: Dashboard, Medicamentos, Control de Inventario, Notificaciones, Reportes, Analytics.",
        ]),
        ("Farmacia", [
            "Acceso limitado al manejo operativo del inventario.",
            "Módulos disponibles: Dashboard, Medicamentos (CRUD completo) y Búsqueda de Medicamentos.",
            "Sin acceso a: Usuarios, Historial, Reportes avanzados ni Analytics.",
        ]),
        ("Jefatura", [
            "Acceso orientado a la supervisión y toma de decisiones.",
            "Módulos disponibles: Dashboard Analítico, Control de Inventario, Notificaciones, Reportes (general, ingresos, rotación) y Analytics.",
            "Sin acceso a: Gestión de Usuarios, Historial de Acciones ni módulo de registro/edición de medicamentos.",
        ]),
    ]

    for perfil_name, perfil_items in perfiles:
        add_heading(doc, f"Perfil: {perfil_name}", 3)
        for item in perfil_items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(item)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    page_break(doc)

    # ── IV. DIAGRAMA DE PAQUETES ──────────────────────────────────────────────
    add_heading(doc, "IV. DIAGRAMA DE PAQUETES", 1)
    add_body(doc, (
        "El siguiente diagrama de paquetes muestra la organización funcional del sistema MOPGIMED, "
        "agrupando los casos de uso según el área funcional que cubren y las dependencias entre paquetes."
    ))
    add_code_block(doc, PUML_PAQUETES_SRS)
    add_caption(doc, "Figura 1: Diagrama de Paquetes — MOPGIMED (Vista Funcional SRS)")

    page_break(doc)

    # ── V. DIAGRAMA DE CASOS DE USO ───────────────────────────────────────────
    add_heading(doc, "V. DIAGRAMA DE CASOS DE USO", 1)
    add_body(doc, (
        "El diagrama de casos de uso presenta los 11 casos de uso consolidados del sistema MOPGIMED, "
        "los actores que interactúan con cada uno y los paquetes funcionales que los agrupan."
    ))
    add_code_block(doc, PUML_CASOS_USO)
    add_caption(doc, "Figura 2: Diagrama de Casos de Uso — MOPGIMED")

    page_break(doc)

    # ── VI. ESCENARIOS DE CASO DE USO ─────────────────────────────────────────
    add_heading(doc, "VI. ESCENARIOS DE CASO DE USO (Narrativas)", 1)

    ui_texts = {
        "UC-001": (
            "La pantalla de inicio de sesión de MOPGIMED presenta un diseño minimalista sobre fondo blanco "
            "con una tarjeta centrada de bordes suaves. En la parte superior de la tarjeta se muestra el "
            "logotipo del sistema: un ícono de triángulo de advertencia (AlertTriangle) de color rojo "
            "(#B91C1C) de 34 píxeles de tamaño, acompañado del título 'MOPGIMED' en negrita y un subtítulo "
            "descriptivo. Debajo se presentan dos campos de texto: 'Usuario' y 'Contraseña', ambos con la "
            "clase form-input que les da un borde gris y un relleno interior cómodo. El campo de contraseña "
            "incluye un botón de ojo (eye-toggle) en el extremo derecho que permite alternar la visibilidad "
            "del texto ingresado. El botón principal 'Iniciar Sesión' ocupa el ancho completo de la tarjeta, "
            "tiene fondo rojo (#B91C1C), texto blanco en negrita y esquinas redondeadas. En la parte inferior "
            "de la tarjeta se muestra una caja de sugerencia con tres credenciales de ejemplo para facilitar "
            "las pruebas del sistema. Si las credenciales son incorrectas, aparece un recuadro de error con "
            "fondo rojo claro y texto rojo oscuro indicando el problema."
        ),
        "UC-002": (
            "El módulo de Gestión de Usuarios presenta en la parte superior un encabezado (div section-header) "
            "con el título 'Gestión de Usuarios' en h1 con la clase page-title (texto oscuro, tamaño grande) "
            "y el subtítulo 'Administración de usuarios y permisos del sistema' en un párrafo con clase "
            "page-subtitle (texto gris claro). A la derecha del encabezado se ubica el botón 'Nuevo Usuario' "
            "con clases btn btn-primary (fondo azul, texto blanco) y un ícono de signo más (Plus icon) a la "
            "izquierda del texto. La tabla principal muestra las columnas: Usuario (nombre en negrita y username "
            "en gris debajo), Email, Rol (badge de color según el rol: rojo para Admin, azul para Farmacia, "
            "verde para Jefatura), Estado (badge verde para Activo, gris para Inactivo), Fecha Creación y "
            "Acciones (ícono lápiz para editar y ícono UserX rojo para desactivar o verde para activar). "
            "Al hacer clic en Nuevo Usuario o en Editar, se abre un modal con fondo oscuro semitransparente "
            "(dark backdrop) y una tarjeta centrada con el formulario de campos: nombre, username, contraseña "
            "(opcional en edición), selector de rol y email. Debajo de la tabla de usuarios se encuentra la "
            "sección 'Historial de Acciones' (indicada con ↗ y la clase de tarjeta card) con su propia tabla: "
            "Fecha/Hora, Usuario, Acción, Módulo (badge de color según el módulo) y Detalle."
        ),
        "UC-011": (
            "El Historial de Acciones se muestra en la sección inferior del módulo Gestión de Usuarios, "
            "separado de la tabla principal de usuarios. Está contenido en una tarjeta (card) con un encabezado "
            "que incluye el texto '↗ Historial de Acciones'. La tabla del historial presenta las columnas: "
            "Fecha/Hora (con formato DD/MM/YYYY HH:MM), Usuario (nombre del usuario que realizó la acción), "
            "Acción (descripción de la operación realizada), Módulo (badge con color diferenciado según el "
            "módulo: rojo para Acceso, azul para Usuarios, verde para Medicamentos, etc.) y Detalle "
            "(información adicional sobre el objeto afectado). Los registros están ordenados del más reciente "
            "al más antiguo. La sección se carga automáticamente al acceder al módulo de Usuarios y no "
            "requiere ninguna acción adicional del Administrador."
        ),
    }

    fig_num = 3
    for uc in UCS:
        add_heading(doc, f"{uc['id']}: {uc['nombre']}", 2)

        # Tabla de UC info
        uc_table = doc.add_table(rows=7, cols=2)
        uc_table.style = 'Table Grid'
        uc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        uc_fields = [
            ("ID del Caso de Uso", uc["id"]),
            ("Nombre", uc["nombre"]),
            ("Actores", uc["actores"]),
            ("Descripción", uc["descripcion"]),
            ("RF Asociado", uc["rf"]),
            ("Precondición", uc["precondicion"]),
            ("Postcondición", uc["postcondicion"]),
        ]
        for i, (campo, detalle) in enumerate(uc_fields):
            row = uc_table.rows[i]
            set_bg(row.cells[0], C_GRIS_CLARO)
            cell_write(row.cells[0], campo, bold=True, size=9)
            cell_write(row.cells[1], detalle, size=9)
        doc.add_paragraph()

        # UI detail
        if uc.get("ui_detail") and uc["id"] in ui_texts:
            add_heading(doc, "Descripción de la Interfaz de Usuario", 3)
            add_body(doc, ui_texts[uc["id"]])

        # Flujo principal
        add_heading(doc, "Flujo Principal", 3)
        for step_num, step in enumerate(uc["flujo_principal"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{step_num}. {step}")
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

        # Flujos alternos
        if uc["flujos_alternos"]:
            add_heading(doc, "Flujos Alternos", 3)
            for alt_name, alt_desc in uc["flujos_alternos"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"• {alt_name}: ")
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run2 = p.add_run(alt_desc)
                run2.font.name = 'Calibri'
                run2.font.size = Pt(10)

        doc.add_paragraph()

    page_break(doc)

    # ── VII. ANÁLISIS DE OBJETOS ──────────────────────────────────────────────
    add_heading(doc, "VII. ANÁLISIS DE OBJETOS (Diagramas de Robustez)", 1)
    add_body(doc, (
        "Los diagramas de robustez presentan el análisis de objetos para los casos de uso principales "
        "del sistema MOPGIMED, utilizando la notación de Frontera (pantallas), Control (lógica interna) "
        "y Entidad (datos almacenados)."
    ))

    robustez_data = [
        ("CU001", "Iniciar Sesión", PUML_ROBUSTEZ_CU001, fig_num),
        ("CU002", "Gestionar Usuarios y Asignar Rol", PUML_ROBUSTEZ_CU002, fig_num + 1),
        ("CU011", "Visualizar Historial de Acciones", PUML_ROBUSTEZ_CU011, fig_num + 2),
    ]
    fig_num += 3

    for rob_id, rob_name, rob_puml, rob_fig in robustez_data:
        add_heading(doc, f"Análisis de Objetos — {rob_id}: {rob_name}", 2)
        add_code_block(doc, rob_puml)
        add_caption(doc, f"Figura {rob_fig}: Análisis de Objetos del Caso de Uso: {rob_name}")
        add_body(doc, (
            f"El diagrama de robustez del {rob_id} muestra cómo el sistema procesa las solicitudes "
            f"del actor a través de las pantallas (Frontera), los procesos internos (Control) y los "
            f"datos almacenados (Entidad), garantizando la trazabilidad completa de cada acción."
        ))
        doc.add_paragraph()

    page_break(doc)

    # ── VIII. DIAGRAMAS DE SECUENCIA ──────────────────────────────────────────
    add_heading(doc, "VIII. DIAGRAMAS DE SECUENCIA", 1)
    add_body(doc, (
        "Los diagramas de secuencia detallan el flujo cronológico de mensajes entre los componentes "
        "del sistema para los casos de uso principales de MOPGIMED."
    ))

    secuencia_data = [
        ("CU001", "Iniciar Sesión", PUML_SECUENCIA_CU001, fig_num),
        ("CU002", "Gestionar Usuarios y Asignar Rol", PUML_SECUENCIA_CU002, fig_num + 1),
        ("CU011", "Visualizar Historial de Acciones", PUML_SECUENCIA_CU011, fig_num + 2),
    ]

    for seq_id, seq_name, seq_puml, seq_fig in secuencia_data:
        add_heading(doc, f"Diagrama de Secuencia — {seq_id}: {seq_name}", 2)
        add_code_block(doc, seq_puml)
        add_caption(doc, f"Figura {seq_fig}: Diagrama de Secuencia del Caso de Uso: {seq_name}")
        doc.add_paragraph()

    # ── GUARDAR ───────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "FD03-SRS-MOPGIMED.docx")
    doc.save(output_path)
    print(f"[OK] Documento generado: {output_path}")


if __name__ == "__main__":
    from generar_srs_epis import build
    build()
