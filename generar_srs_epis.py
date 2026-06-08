"""
Genera FD03-SRS-MOPGIMED.docx copiando la estructura de FD03-EPIS y actualizando:
  - Cuadro de Requerimientos Funcionales (17 RF, tabla única)
  - Escenarios de Caso de Uso (narrativas según diagrama de casos de uso)
"""

import os
import shutil
import re
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from uc_narrativas_mopgimed import NARRATIVAS
from trazabilidad_mopgimed import TRAZABILIDAD_SPRINT, INVENTARIO_IMPLEMENTACION

C_ROJO        = 'B91C1C'
C_AZUL_HDR    = '1E3A5F'
C_AZUL_RF     = 'DBEAFE'
C_ROJO_CLARO  = 'FEF2F2'
C_GRIS_CLARO  = 'F1F5F9'
C_BLANCO      = 'FFFFFF'
C_OSCURO      = '1E293B'


def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_rich_text(paragraph, text, bold_default=False, font_size=10, font_color=None):
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        is_bold = part.startswith('**') and part.endswith('**')
        content = part[2:-2] if is_bold else part
        if content:
            run = paragraph.add_run(content)
            run.bold = is_bold or bold_default
            run.font.size = Pt(font_size)
            run.font.name = 'Calibri'
            if font_color:
                run.font.color.rgb = hex_to_rgb(font_color)


def _set_table_style(tbl):
    try:
        tbl.style = 'Table Grid'
    except KeyError:
        try:
            tbl.style = 'Normal Table'
        except KeyError:
            pass


def cell_write(cell, text, bold=False, size=9, color=None, center=False):
    set_cell_margins(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(str(text))
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = hex_to_rgb(color)


def _apply_col_widths(tbl, widths_cm):
    try:
        tbl.autofit = False
    except Exception:
        pass
    for row in tbl.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def build_info_table(doc, fields):
    """Tabla estilo EPIS: filas Campo | Detalle (2 columnas)."""
    tbl = doc.add_table(rows=len(fields), cols=2)
    _set_table_style(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, (campo, detalle) in enumerate(fields):
        c0, c1 = tbl.rows[r_idx].cells[0], tbl.rows[r_idx].cells[1]
        set_cell_margins(c0)
        set_cell_margins(c1)
        set_cell_bg(c0, C_GRIS_CLARO)
        add_rich_text(c0.paragraphs[0], campo, bold_default=True, font_size=10)
        add_rich_text(c1.paragraphs[0], detalle, font_size=10)
    for row in tbl.rows:
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(11.3)
    doc.add_paragraph()


def build_events_table(doc, titulo, pasos):
    """Tabla Curso normal: fila título, encabezado Usuario|Sistema, pasos numerados."""
    rows = [
        [titulo, titulo],
        ['Usuario', 'Sistema'],
    ]
    for i, (u, s) in enumerate(pasos, 1):
        u_txt = f'{i}. {u}' if u else ''
        s_txt = f'{i}. {s}' if s else ''
        rows.append([u_txt, s_txt])

    tbl = doc.add_table(rows=len(rows), cols=2)
    _set_table_style(tbl)
    for r_idx, row_data in enumerate(rows):
        for c_idx, txt in enumerate(row_data):
            cell = tbl.rows[r_idx].cells[c_idx]
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            if r_idx == 0:
                set_cell_bg(cell, C_GRIS_CLARO)
                add_rich_text(para, txt, bold_default=True, font_size=10)
            elif r_idx == 1:
                set_cell_bg(cell, C_ROJO)
                add_rich_text(para, txt, bold_default=True, font_size=10, font_color=C_BLANCO)
            else:
                if r_idx % 2 == 0:
                    set_cell_bg(cell, C_ROJO_CLARO)
                add_rich_text(para, txt, font_size=10)
    for row in tbl.rows:
        row.cells[0].width = Cm(7.25)
        row.cells[1].width = Cm(7.25)
    doc.add_paragraph()


def add_uc_narrative(doc, uc):
    p = doc.add_paragraph()
    run = p.add_run(f"{uc['id']} – {uc['nombre']}")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = hex_to_rgb(C_ROJO)

    fields = [
        ('Caso de uso', f"{uc['id']} – {uc['nombre']}"),
        ('Actores', uc['actores']),
        ('Propósito', uc['proposito']),
        ('Tipo', 'Obligatorio (X) / Opcional ( )'),
        ('Requisito ID (RF)', uc['rf']),
        ('Versión', '1.0'),
        ('Descripción', uc['descripcion']),
        ('Precondición', uc['precondicion']),
        ('Postcondición', uc['postcondicion']),
    ]
    build_info_table(doc, fields)

    for curso in uc['cursos']:
        build_events_table(doc, curso['titulo'], curso['pasos'])

    for alt_titulo, alt_pasos in uc['alternos']:
        build_events_table(doc, alt_titulo, alt_pasos)


# 17 RF — tabla única (sin agrupación por rol)
RF_ROWS = [
    ("RF-001", "Iniciar sesión",
     "Permite validar las credenciales del usuario y autorizar el acceso al sistema según el rol asignado.",
     "RF", "Alta", "Punto de entrada obligatorio; sin autenticación no hay acceso protegido.", "Salas Jiménez, W."),
    ("RF-002", "Gestionar usuarios y asignar rol",
     "Permite registrar, modificar, habilitar o deshabilitar usuarios y asignar perfiles de acceso desde el formulario de gestión.",
     "RF", "Alta", "Control de accesos y permisos del sistema.", "Vargas G., A. / Salas J., W."),
    ("RF-003", "Visualizar historial de acciones",
     "Permite consultar el registro de actividades de los usuarios para auditoría interna.",
     "RF", "Media", "Trazabilidad sin afectar la operación diaria del inventario.", "Salas Jiménez, W."),
    ("RF-004", "Registrar lote de medicamento existente",
     "Permite registrar lote, fecha de vencimiento, unidad y cantidad por envase de un medicamento ya catalogado.",
     "RF", "Alta", "Metadatos operativos del ingreso físico al almacén.", "Vargas Gutierrez, A."),
    ("RF-005", "Registrar medicamento nuevo con asistencia IA",
     "Permite buscar con IA, seleccionar y registrar un medicamento nuevo con metadatos de lote.",
     "RF", "Alta", "Alta de productos al catálogo con estandarización de datos.", "Vargas Gutierrez, A."),
    ("RF-006", "Consultar medicamento (disponibilidad y precio)",
     "Permite buscar y visualizar stock y precio de venta sin modificar el inventario.",
     "RF", "Alta", "Consulta operativa en caja y farmacia.", "Vargas Gutierrez, A."),
    ("RF-007", "Listar medicamentos (buscador + exportar Excel)",
     "Permite visualizar el inventario completo, buscar y exportar el listado a Excel.",
     "RF", "Alta", "Vista principal del catálogo para supervisión.", "Salas Jiménez, W."),
    ("RF-008", "Filtrar listado de medicamentos",
     "Permite refinar el listado por laboratorio u otros criterios de búsqueda.",
     "RF", "Alta", "Localización rápida de productos por fabricante.", "Salas Jiménez, W."),
    ("RF-009", "Modificar medicamento",
     "Permite corregir o actualizar la información de un medicamento registrado.",
     "RF", "Alta", "Mantenimiento del catálogo ante cambios o errores.", "Vargas Gutierrez, A."),
    ("RF-010", "Visualizar semáforo de stock de medicamentos",
     "Permite mostrar indicadores agregados Normal, Bajo y Crítico del inventario.",
     "RF", "Alta", "Visión rápida del estado global del stock.", "Salas Jiménez, W."),
    ("RF-011", "Visualizar semáforo de medicamentos por vencer / vencidos",
     "Permite identificar productos próximos a vencer o ya vencidos.",
     "RF", "Alta", "Prevención de pérdidas por caducidad.", "Salas Jiménez, W."),
    ("RF-012", "Listar medicamentos con semáforo",
     "Permite listar el inventario mostrando el estado de stock por medicamento.",
     "RF", "Alta", "Detalle operativo con código de colores por fila.", "Salas Jiménez, W."),
    ("RF-013", "Filtrar lista de vencimientos",
     "Permite aplicar filtros por laboratorio, mes o año sobre medicamentos por vencer.",
     "RF", "Media", "Consulta precisa de vencimientos.", "Salas Jiménez, W."),
    ("RF-014", "Reporte general de medicamentos (filtros + exportar Excel)",
     "Permite emitir reporte consolidado del inventario con filtros y exportación.",
     "RF", "Alta", "Supervisión gerencial del inventario.", "Vargas Gutierrez, A."),
    ("RF-015", "Reporte de ingresos y salidas (filtro por tiempo)",
     "Permite consultar movimientos históricos de entradas y salidas por periodo.",
     "RF", "Alta", "Análisis de flujo del inventario.", "Vargas Gutierrez, A."),
    ("RF-016", "Reporte de rotación de medicamentos",
     "Permite identificar medicamentos de mayor rotación y generar su reporte.",
     "RF", "Media", "Apoyo a decisiones de compra.", "Salas Jiménez, W."),
    ("RF-017", "Visualizar dashboard analítico (operativo + inteligente ML/IA)",
     "Permite mostrar indicadores, gráficos y análisis predictivo de demanda y tendencias.",
     "RF", "Alta", "Toma de decisiones estratégicas con datos e IA.", "Vargas Gutierrez, A."),
]


def _paragraph_text(p):
    return ''.join(r.text for r in p.runs).strip()


def _find_paragraph_index(doc, contains):
    for i, p in enumerate(doc.paragraphs):
        if contains in _paragraph_text(p):
            return i
    return -1


def _remove_body_elements_between(doc, start_substr, end_substr, skip_start=True):
    """Elimina párrafos y tablas entre dos encabezados."""
    body = doc.element.body
    children = list(body)
    removing = False
    to_remove = []

    for el in children:
        tag = el.tag.split('}')[-1]
        text = ''
        if tag == 'p':
            texts = [t.text for t in el.iter() if t.text]
            text = ''.join(texts).strip()
        elif tag == 'tbl':
            text = '__TABLE__'

        if not removing and start_substr in text:
            removing = True
            if skip_start:
                continue
        if removing and end_substr in text:
            break
        if removing:
            to_remove.append(el)

    for el in to_remove:
        body.remove(el)


def _replace_rf_table(doc):
    """Reemplaza tabla RF con formato legible y anchos de columna fijos."""
    tbl = doc.tables[2]
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[-1]._tr)

    headers = ['ID', 'Nombre', 'Descripción', 'Tipo', 'Prioridad',
               'Justificación de Prioridad', 'Task Owner']
    # Total ~17 cm (A4 con márgenes 2.5 cm)
    widths = [1.3, 2.8, 4.8, 0.9, 1.1, 3.8, 2.3]

    for i, h in enumerate(headers):
        if i < len(tbl.rows[0].cells):
            set_cell_bg(tbl.rows[0].cells[i], C_AZUL_HDR)
            cell_write(tbl.rows[0].cells[i], h, bold=True, size=8, color=C_BLANCO, center=True)

    for row_data in RF_ROWS:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            if i < len(row.cells):
                set_cell_bg(row.cells[i], C_AZUL_RF)
                cell_write(row.cells[i], val, size=8, center=(i in (0, 3, 4)))

    _apply_col_widths(tbl, widths)


def _build_data_table(doc, headers, rows, widths_cm, header_color=C_AZUL_HDR, row_color=C_AZUL_RF):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _set_table_style(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        set_cell_bg(tbl.rows[0].cells[i], header_color)
        cell_write(tbl.rows[0].cells[i], h, bold=True, size=8, color=C_BLANCO, center=True)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            set_cell_bg(tbl.rows[r_idx + 1].cells[c_idx], row_color)
            cell_write(tbl.rows[r_idx + 1].cells[c_idx], val, size=8,
                       center=(c_idx in (0, 3, 4) if len(headers) > 4 else c_idx == 0))
    _apply_col_widths(tbl, widths_cm)
    doc.add_paragraph()
    return tbl


def _add_heading_paragraph(doc, text, level=3):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11 if level == 3 else 12)
    run.font.color.rgb = hex_to_rgb(C_ROJO if level == 3 else C_OSCURO)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def _build_traceability_sections(doc):
    """Punto 1 y 3 de trazabilidad requerida por el sprint."""
    _add_heading_paragraph(doc, 'Requerimientos cubiertos por el sprint', level=3)
    _add_body_paragraph(doc, (
        'La siguiente matriz relaciona cada requerimiento funcional (RF) del presente SRS con el '
        'caso de uso (UC), el sprint de desarrollo planificado, el estado de implementación en '
        'MOPGIMED y la referencia a las vistas del Documento de Arquitectura de Software (SAD). '
        'Esta sección responde al ítem de trazabilidad: Documento SRS — requerimientos cubiertos por el sprint.'
    ))
    sprint_rows = [
        (rf, nombre, uc, f'Sprint {sp}', estado, sad)
        for rf, nombre, uc, sp, estado, sad in TRAZABILIDAD_SPRINT
    ]
    _build_data_table(
        doc,
        ['RF', 'Requerimiento', 'Caso de uso', 'Sprint', 'Estado', 'Referencia SAD'],
        sprint_rows,
        [1.1, 2.6, 1.4, 1.0, 1.2, 5.7],
    )

    _add_heading_paragraph(doc, 'Inventario de componentes de implementación', level=3)
    _add_body_paragraph(doc, (
        'El inventario siguiente identifica los componentes creados en la implementación del sistema '
        '(frontend React, backend Express y base de datos MySQL), indicando su trazabilidad con los '
        'requerimientos del SRS y los paquetes/vistas descritos en el SAD. '
        'Responde al ítem: Implementación — inventario de componentes que responden al SRS y SAD.'
    ))
    impl_rows = [
        (comp, tipo, ruta, rf, uc, sad)
        for comp, tipo, ruta, rf, uc, sad in INVENTARIO_IMPLEMENTACION
    ]
    _build_data_table(
        doc,
        ['Componente', 'Tipo', 'Ruta / Archivo', 'RF', 'UC', 'Referencia SAD'],
        impl_rows,
        [2.2, 1.3, 3.8, 1.5, 1.5, 3.7],
    )


def _insert_before_heading(doc, heading_substr, builder_fn):
    anchor_el = None
    for el in doc.element.body:
        if el.tag.split('}')[-1] != 'p':
            continue
        text = ''.join(t.text for t in el.iter() if t.text).strip()
        if heading_substr in text:
            anchor_el = el
            break
    if anchor_el is None:
        builder_fn(doc)
        return
    tmp = Document()
    tmp._part = doc._part
    builder_fn(tmp)
    for el in reversed(list(tmp.element.body)):
        tag = el.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            anchor_el.addprevious(deepcopy(el))


def _find_modelo_logico_element(doc):
    """Localiza el párrafo Modelo Lógico (sección de diseño, no el título del proyecto)."""
    after_escenarios = False
    for el in doc.element.body:
        if el.tag.split('}')[-1] != 'p':
            continue
        text = ''.join(t.text for t in el.iter() if t.text).strip()
        if 'Escenarios de Caso de Uso' in text:
            after_escenarios = True
            continue
        if after_escenarios and 'Modelo L' in text and 'gico' in text:
            return el
    return None


def _insert_elements_before_el(anchor_el, elements):
    for el in elements:
        anchor_el.addprevious(deepcopy(el))


def _remove_duplicate_tail_narratives(doc):
    """Elimina bloque duplicado Escenarios+narrativas si quedó al final del documento."""
    body = doc.element.body
    children = list(body)
    esc_indices = []
    for i, el in enumerate(children):
        if el.tag.split('}')[-1] != 'p':
            continue
        text = ''.join(t.text for t in el.iter() if t.text).strip()
        if 'Escenarios de Caso de Uso' in text and 'narrativa' in text.lower():
            esc_indices.append(i)
    if len(esc_indices) < 2:
        return
    start = esc_indices[1]
    end = len(children)
    for i in range(start, len(children)):
        el = children[i]
        if el.tag.split('}')[-1] == 'p':
            text = ''.join(t.text for t in el.iter() if t.text).strip()
            if text.startswith('CONCLUSIONES'):
                end = i
                break
    for el in children[start:end]:
        body.remove(el)


def _build_narrative_elements(source_doc):
    """Genera elementos XML de narrativas usando estilos del documento EPIS."""
    tmp = Document()
    tmp._part = source_doc._part
    for uc in NARRATIVAS:
        add_uc_narrative(tmp, uc)
    out = []
    for el in tmp.element.body:
        tag = el.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            out.append(el)
    return out


def build():
    root = os.path.dirname(os.path.abspath(__file__))
    epis_name = next(f for f in os.listdir(root) if f.startswith('FD03-EPIS'))
    epis_path = os.path.join(root, epis_name)
    out_path = os.path.join(root, 'FD03-SRS-MOPGIMED.docx')

    try:
        shutil.copy2(epis_path, out_path)
    except PermissionError:
        out_path = os.path.join(root, 'FD03-SRS-MOPGIMED-GENERADO.docx')
        shutil.copy2(epis_path, out_path)
        print('[AVISO] El archivo SRS estaba abierto. Guardado como:', out_path)
    doc = Document(out_path)

    for p in doc.paragraphs[:35]:
        t = _paragraph_text(p)
        if 'Documento de Especificación' in t or 'Documento de Especificaci' in t:
            for r in p.runs:
                if 'Requerimientos' in r.text:
                    r.text = 'SRS — Especificación de Requerimientos de Software (MOPGIMED)'

    _replace_rf_table(doc)
    _insert_before_heading(doc, 'Reglas de Negocio', _build_traceability_sections)

    _remove_body_elements_between(
        doc,
        'Escenarios de Caso de Uso (narrativa)',
        'Modelo Lógico',
    )

    modelo_el = _find_modelo_logico_element(doc)
    narrative_els = _build_narrative_elements(doc)

    if modelo_el is not None and narrative_els:
        _insert_elements_before_el(modelo_el, narrative_els)
    else:
        for uc in NARRATIVAS:
            add_uc_narrative(doc, uc)

    _remove_duplicate_tail_narratives(doc)

    doc.save(out_path)
    print(f'[OK] Generado: {out_path}')
    print(f'     RF funcionales: {len(RF_ROWS)}')
    print(f'     Narrativas UC:  {len(NARRATIVAS)}')


if __name__ == '__main__':
    build()
