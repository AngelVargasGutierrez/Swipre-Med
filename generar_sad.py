"""
generar_sad.py
Genera FD04-SAD-MOPGIMED.docx — Documento de Arquitectura de Software
Sistema Inteligente de Inventario Farmacéutico — MOPGIMED
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
#  COLORES
# ─────────────────────────────────────────────────────────────────────────────
C_ROJO         = 'B91C1C'
C_OSCURO       = '1E293B'
C_AZUL_OSCURO  = '1E3A5F'
C_AZUL_CLARO   = 'DBEAFE'
C_VERDE_CLARO  = 'DCFCE7'
C_GRIS_CLARO   = 'F1F5F9'
C_GRIS_CODE    = 'F1F5F9'
C_BLANCO       = 'FFFFFF'
C_NEGRO        = '111827'
C_AMARILLO     = 'FEF9C3'

UNIV     = "Universidad Privada de Tacna"
FACULTAD = "Facultad de Ingeniería"
ESCUELA  = "Escuela Profesional de Ingeniería de Sistemas"
CURSO    = "Ingeniería de Software I — SI885"
DOCENTE  = "Mg. Ing. Rosa María Flores Condori"
YEAR     = "2025"
TEAM_MEMBERS = [
    "Vargas Gutierrez, Angel",
    "Lopez Mamani, Carlos",
    "Quispe Flores, Maria",
    "Condori Tapia, Luis",
]


# ─────────────────────────────────────────────────────────────────────────────
#  UTILIDADES
# ─────────────────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_bg(cell, color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color.lstrip('#'))
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def cell_write(cell, text, bold=False, size=10, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p   = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = hex_to_rgb(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size      = Pt(14)
        run.font.color.rgb = hex_to_rgb(C_ROJO)
    elif level == 2:
        run.font.size      = Pt(12)
        run.font.color.rgb = hex_to_rgb(C_OSCURO)
    else:
        run.font.size      = Pt(11)
        run.font.color.rgb = hex_to_rgb(C_OSCURO)
    return p


def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = hex_to_rgb(C_NEGRO)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold       = True
    run.italic     = True
    run.font.name  = 'Calibri'
    run.font.size  = Pt(9)
    run.font.color.rgb = hex_to_rgb(C_OSCURO)


def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_bg(cell, C_GRIS_CODE)
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    '100')
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7)
    run.font.color.rgb = hex_to_rgb('1E293B')
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def set_margins(doc, cm_val=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(cm_val)
        section.bottom_margin = Cm(cm_val)
        section.left_margin   = Cm(cm_val)
        section.right_margin  = Cm(cm_val)


# ─────────────────────────────────────────────────────────────────────────────
#  PLANTUMLS
# ─────────────────────────────────────────────────────────────────────────────

PUML_PAQUETES_SAD = """\
@startuml MOPGIMED Diagrama de Paquetes SAD
top to bottom direction
title "MOPGIMED — Diagrama de Paquetes (Vista Tecnica)\\nDocumento de Arquitectura de Software (SAD)"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  PackageBackgroundColor #F8FAFC  PackageBorderColor #1E293B
  PackageBorderThickness 2  PackageFontStyle Bold
  ClassBackgroundColor #FFFFFF  ClassBorderColor #94A3B8
  ClassHeaderBackgroundColor #1E3A5F  ClassFontColor #EEEEEE
  ArrowColor #B91C1C  ArrowThickness 1.5
}
package "Frontend  —  src/" as FRONTEND #EFF6FF {
  package "src/pages/" as PAGES {
    class "Login.jsx" <<page>>
    class "Dashboard.jsx" <<page>>
    class "Medicamentos.jsx" <<page>>
    class "NuevoMedicamento.jsx" <<page>>
    class "ControlInventario.jsx" <<page>>
    class "Notificaciones.jsx" <<page>>
    class "Analytics.jsx" <<page>>
    class "Reportes.jsx" <<page>>
    class "Usuarios.jsx" <<page>>
  }
  package "src/components/" as COMPONENTS {
    class "Sidebar.jsx" <<component>>
    class "Loader.jsx" <<component>>
  }
  package "src/context/" as CONTEXT {
    class "AuthContext.jsx" <<context>>
  }
  package "src/services/" as SERVICES {
    class "api.js" <<service>>
  }
}
package "Backend  —  backend/" as BACKEND #FFFBEB {
  package "backend/routes/" as ROUTES {
    class "auth.js" <<route>>
    class "users.js" <<route>>
    class "medicamentos.js" <<route>>
    class "inventario.js" <<route>>
    class "notificaciones.js" <<route>>
    class "reportes.js" <<route>>
    class "dashboard.js" <<route>>
    class "analytics.js" <<route>>
  }
  package "backend/controllers/" as CONTROLLERS {
    class "authController.js" <<controller>>
    class "userController.js" <<controller>>
    class "medicamentoController.js" <<controller>>
    class "inventarioController.js" <<controller>>
    class "notificacionController.js" <<controller>>
    class "reporteController.js" <<controller>>
    class "dashboardController.js" <<controller>>
    class "analyticsController.js" <<controller>>
  }
  package "backend/models/" as MODELS {
    class "userModel.js" <<model>>
    class "medicamentoModel.js" <<model>>
    class "inventarioModel.js" <<model>>
    class "notificacionModel.js" <<model>>
    class "reporteModel.js" <<model>>
    class "dashboardModel.js" <<model>>
    class "analyticsModel.js" <<model>>
  }
  package "backend/db/" as DB_CONN {
    class "connection.js" <<database>>
    class "index.js" <<entrypoint>>
  }
}
package "Base de Datos  —  MySQL: swipre_med" as DATABASE #ECFDF5 {
  package "Tablas Principales" {
    class "users" <<table>>
    class "medicamentos" <<table>>
    class "notificaciones" <<table>>
  }
  package "Tablas de Auditoria" {
    class "historial_acciones" <<table>>
    class "alertas" <<table>>
  }
  package "Tablas de Analisis" {
    class "movimientos_semana" <<table>>
    class "comparativa_mensual" <<table>>
    class "prediccion_demanda" <<table>>
    class "tendencias" <<table>>
    class "predicciones_criticas" <<table>>
  }
}
SERVICES  ..> ROUTES      : HTTP REST (Axios)
ROUTES    ..> CONTROLLERS : delega logica de negocio
CONTROLLERS ..> MODELS    : accede a los datos
MODELS    ..> DB_CONN     : ejecuta consultas SQL
DB_CONN   ..> DATABASE    : conexion MySQL2 pool
@enduml"""

PUML_CASOS_USO = """\
@startuml DiagramaCasosDeUso MOPGIMED
title "Diagrama de Casos de Uso — MOPGIMED\\nSistema Inteligente de Inventario Farmacéutico"
left to right direction
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam ArrowColor #B91C1C
skinparam ArrowThickness 1.2
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam usecase { BorderColor #475569  BackgroundColor #FFFFFF  FontColor #1E293B }
actor "Administrador" as Admin
actor "Farmacia" as Farm
actor "Jefatura" as Jef
package "Acceso al Sistema" #EFF6FF {
  usecase "UC-001: Iniciar Sesion" as UC01
}
package "Administracion de Usuarios" #FFFBEB {
  usecase "UC-002: Gestionar Usuarios\\ny Asignar Rol" as UC02
  usecase "UC-011: Visualizar Historial\\nde Acciones" as UC11
}
package "Gestion de Medicamentos" #ECFDF5 {
  usecase "UC-003: Gestion de\\nMedicamentos - CRUD" as UC03
  usecase "UC-004: Busqueda y Filtrado\\nde Medicamentos" as UC04
}
package "Control de Inventario" #FFF7ED {
  usecase "UC-005: Control de Inventario" as UC05
}
package "Alertas y Notificaciones" #FFF1F2 {
  usecase "UC-006: Gestion de Notificaciones" as UC06
}
package "Reportes" #F5F3FF {
  usecase "UC-007: Reporte General" as UC07
  usecase "UC-008: Reporte Ingresos y Salidas" as UC08
  usecase "UC-009: Reporte Rotacion" as UC09
}
package "Dashboard Analitico" #F0FDF4 {
  usecase "UC-010: Dashboard Analitico" as UC10
}
Admin --> UC01
Admin --> UC02
Admin --> UC11
Admin --> UC03
Admin --> UC04
Admin --> UC05
Admin --> UC06
Admin --> UC07
Admin --> UC08
Admin --> UC09
Admin --> UC10
Farm --> UC01
Farm --> UC03
Farm --> UC04
Jef --> UC01
Jef --> UC05
Jef --> UC06
Jef --> UC07
Jef --> UC08
Jef --> UC09
Jef --> UC10
@enduml"""

# Robustez SAD — nombres técnicos
PUML_ROB_SAD_CU001 = """\
@startuml Robustez SAD CU001 MOPGIMED
top to bottom direction
title "Diagrama de Robustez SAD — CU001: Iniciar Sesion\\nMOPGIMED · Nombres Tecnicos"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ArrowColor #B91C1C  ArrowThickness 1.5
}
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam boundary { BorderColor #1D4ED8  BackgroundColor #EFF6FF  FontColor #1E3A8A }
skinparam control { BorderColor #B45309  BackgroundColor #FFFBEB  FontColor #78350F }
skinparam entity { BorderColor #065F46  BackgroundColor #ECFDF5  FontColor #064E3B }
skinparam package { BorderColor #CBD5E1  BackgroundColor #F8FAFC  FontStyle Bold }
actor "Usuario del Sistema" as ACTOR
package "Entidad" as PKG_E {
  entity "users\\n(userModel.js)\\nusername · password · role · estado" as E1
  entity "Permisos por Rol\\n(AuthContext.jsx)\\nadmin · farmacia · jefatura" as E2
  entity "historial_acciones\\n(reporteModel.js)\\nfecha · usuario · accion · modulo" as E3
  E1 -[hidden]right-> E2
  E2 -[hidden]right-> E3
}
package "Frontera" as PKG_B {
  boundary "Login.jsx\\nPOST /api/auth/login" as B1
  boundary "Error Message div\\n(estado local React)" as B3
  boundary "Dashboard.jsx / App Router\\n(rutas protegidas)" as B2
  B1 -[hidden]down-> B3
  B3 -[hidden]down-> B2
}
package "Control" as PKG_C {
  control "authController.js\\nlogin(req, res)" as C1
  control "authController.js\\ngetMenu(role)" as C2
  control "reporteController.js\\nregistrarAccion()" as C3
  C1 -[hidden]down-> C2
  C2 -[hidden]down-> C3
}
PKG_E -[hidden]down-> PKG_B
PKG_B -[hidden]right-> PKG_C
ACTOR --> B1  : 1. POST {username, password}
B1    --> C1  : 2. api.js → /api/auth/login
C1    --> E1  : 3. userModel.findByCredentials()
C1    --> B3  : 4a. res.status(401) credenciales invalidas
B3    --> ACTOR : muestra mensaje de error
C1    --> C2  : 4b. Token JWT generado
C2    --> E2  : 5. Consulta permisos del rol
C1    --> C3  : 6. Notifica ingreso exitoso
C3    --> E3  : 7. INSERT historial_acciones
C2    --> B2  : 8. Redirige segun rol
B2    --> ACTOR : 9. Panel principal cargado
@enduml"""

PUML_ROB_SAD_CU002 = """\
@startuml Robustez SAD CU002 MOPGIMED
top to bottom direction
title "Diagrama de Robustez SAD — CU002: Gestionar Usuarios\\nMOPGIMED · Nombres Tecnicos"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ArrowColor #B91C1C  ArrowThickness 1.5
}
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam boundary { BorderColor #1D4ED8  BackgroundColor #EFF6FF  FontColor #1E3A8A }
skinparam control { BorderColor #B45309  BackgroundColor #FFFBEB  FontColor #78350F }
skinparam entity { BorderColor #065F46  BackgroundColor #ECFDF5  FontColor #064E3B }
skinparam package { BorderColor #CBD5E1  BackgroundColor #F8FAFC  FontStyle Bold }
actor "Administrador" as ACTOR
package "Entidad" as PKG_E {
  entity "users\\n(userModel.js)\\nid · name · username · role · email · estado" as E1
  entity "Roles definidos\\n(backend/routes/users.js)\\nadmin · farmacia · jefatura" as E2
  entity "historial_acciones\\n(reporteModel.js)\\nfecha · admin · accion · usuario afectado" as E3
  E1 -[hidden]right-> E2
  E2 -[hidden]right-> E3
}
package "Frontera" as PKG_B {
  boundary "Usuarios.jsx\\nGET/POST /api/users" as B1
  boundary "Modal Form (Usuarios.jsx)\\nFormulario nuevo/editar usuario" as B2
  boundary "Error Toast (Usuarios.jsx)\\nestado local React useState" as B3
  B1 -[hidden]down-> B2
  B2 -[hidden]down-> B3
}
package "Control" as PKG_C {
  control "userController.js\\ngetAll() · create() · update() · toggleStatus()" as C1
  control "userController.js\\nvalidateUserData()" as C2
  control "userController.js\\nassignRoleLabel()" as C3
  control "reporteController.js\\nregistrarAccion()" as C4
  C1 -[hidden]down-> C2
  C2 -[hidden]down-> C3
  C3 -[hidden]down-> C4
}
PKG_E -[hidden]down-> PKG_B
PKG_B -[hidden]right-> PKG_C
ACTOR --> B1  : 1. GET /api/users (cargar lista)
B1    --> C1  : 2. userController.getAll()
C1    --> E1  : 3. userModel.findAll()
E1    --> B1  : 4. JSON array usuarios
ACTOR --> B2  : 5. POST/PUT datos del formulario
B2    --> C2  : 6. Validacion de campos
C2    --> B3  : 7a. res.status(400) datos invalidos
B3    --> ACTOR : muestra toast de error
C2    --> C3  : 7b. Datos correctos
C3    --> E2  : 8. Asigna roleLabel al usuario
C3    --> C1  : 9. Procede a guardar
C1    --> E1  : 10. userModel.create() o update()
C1    --> C4  : 11. registrarAccion(admin, accion)
C4    --> E3  : 12. INSERT historial_acciones
C1    --> B1  : 13. Actualiza lista de usuarios
B1    --> ACTOR : 14. Muestra resultado en tabla
@enduml"""

PUML_ROB_SAD_CU011 = """\
@startuml Robustez SAD CU011 MOPGIMED
top to bottom direction
title "Diagrama de Robustez SAD — CU011: Visualizar Historial\\nMOPGIMED · Nombres Tecnicos"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ArrowColor #B91C1C  ArrowThickness 1.5
}
skinparam actor { BorderColor #B91C1C  BackgroundColor #FFF1F2  FontColor #7F1D1D  FontStyle Bold }
skinparam boundary { BorderColor #1D4ED8  BackgroundColor #EFF6FF  FontColor #1E3A8A }
skinparam control { BorderColor #B45309  BackgroundColor #FFFBEB  FontColor #78350F }
skinparam entity { BorderColor #065F46  BackgroundColor #ECFDF5  FontColor #064E3B }
skinparam package { BorderColor #CBD5E1  BackgroundColor #F8FAFC  FontStyle Bold }
actor "Administrador" as ACTOR
package "Entidad" as PKG_E {
  entity "historial_acciones\\nid · fecha · usuario · accion · modulo · detalle" as E1
  entity "users\\n(userModel.js)\\nnombre · role para enriquecer historial" as E2
  E1 -[hidden]right-> E2
}
package "Frontera" as PKG_B {
  boundary "Usuarios.jsx\\n(seccion historial)\\nGET /api/reportes/historial" as B1
  boundary "Tabla Historial (Usuarios.jsx)\\nfecha · usuario · accion · modulo badge · detalle" as B2
  B1 -[hidden]down-> B2
}
package "Control" as PKG_C {
  control "reporteController.js\\ngetHistorial(req, res)" as C1
  control "reporteController.js\\nregistrarAccion(usuario, accion, modulo, detalle)" as C2
  C1 -[hidden]down-> C2
}
PKG_E -[hidden]down-> PKG_B
PKG_B -[hidden]right-> PKG_C
ACTOR --> B1  : 1. Accede a Usuarios.jsx
B1    --> C1  : 2. GET /api/reportes/historial
C1    --> E1  : 3. reporteModel.getHistorial()\\nSELECT * ORDER BY fecha DESC
E1    --> C1  : 4. JSON array de registros
C1    --> B2  : 5. res.json(historial)
B2    --> ACTOR : 6. Renderiza tabla de historial
ACTOR --> B1  : 7. Realiza accion (CRUD usuarios)
B1    --> C2  : 8. Llama registrarAccion()
C2    --> E2  : 9. Consulta nombre del admin
C2    --> E1  : 10. INSERT INTO historial_acciones
E1    --> C2  : 11. Confirma insercion
C2    --> B1  : 12. Accion registrada automaticamente
@enduml"""

PUML_SECUENCIA_CU001 = """\
@startuml Secuencia SAD CU001 MOPGIMED
title "Diagrama de Secuencia (Tecnico) — CU001: Iniciar Sesion\\nMOPGIMED"
autonumber
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam sequenceArrowColor #B91C1C
skinparam sequenceArrowThickness 1.5
actor "Usuario" as ACTOR
box "Frontend (src/)" #EFF6FF
  participant "Login.jsx" as B1
  participant "AuthContext.jsx" as CTX
  participant "api.js" as API
end box
box "Backend (backend/)" #FFFBEB
  participant "routes/auth.js" as ROUTE
  participant "authController.js" as CTRL
  participant "userModel.js" as MODEL
  participant "reporteController.js" as REP
end box
box "Base de Datos" #ECFDF5
  participant "users" as DB_USERS
  participant "historial_acciones" as DB_HIST
end box
ACTOR -> B1 : Ingresa username y password
B1 -> API : api.login({username, password})
API -> ROUTE : POST /api/auth/login
ROUTE -> CTRL : authController.login(req, res)
CTRL -> MODEL : userModel.findByCredentials(username)
MODEL -> DB_USERS : SELECT * FROM users WHERE username=?
DB_USERS --> MODEL : Fila del usuario
MODEL --> CTRL : Objeto usuario o null
alt Credenciales invalidas o estado Inactivo
  CTRL --> API : res.status(401) {error}
  API --> B1 : Promise.reject
  B1 --> ACTOR : Muestra mensaje de error en pantalla
else Credenciales validas y estado Activo
  CTRL -> REP : registrarAccion(nombre, "Inicio sesion", "Acceso", "")
  REP -> DB_HIST : INSERT INTO historial_acciones
  DB_HIST --> REP : OK
  CTRL --> API : res.json({token, user})
  API --> CTX : Guarda token en localStorage
  CTX --> B1 : Actualiza estado de autenticacion
  B1 --> ACTOR : Redirige a Dashboard segun rol
end
@enduml"""

PUML_SECUENCIA_CU002 = """\
@startuml Secuencia SAD CU002 MOPGIMED
title "Diagrama de Secuencia (Tecnico) — CU002: Gestionar Usuarios\\nMOPGIMED"
autonumber
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam sequenceArrowColor #B91C1C
skinparam sequenceArrowThickness 1.5
actor "Administrador" as ACTOR
box "Frontend (src/)" #EFF6FF
  participant "Usuarios.jsx" as PAGE
  participant "api.js" as API
end box
box "Backend (backend/)" #FFFBEB
  participant "routes/users.js" as ROUTE
  participant "userController.js" as CTRL
  participant "userModel.js" as MODEL
  participant "reporteController.js" as REP
end box
box "Base de Datos" #ECFDF5
  participant "users" as DB_USERS
  participant "historial_acciones" as DB_HIST
end box
== Cargar lista de usuarios ==
ACTOR -> PAGE : Navega a /usuarios
PAGE -> API : api.getUsers()
API -> ROUTE : GET /api/users
ROUTE -> CTRL : userController.getAll(req, res)
CTRL -> MODEL : userModel.findAll()
MODEL -> DB_USERS : SELECT * FROM users ORDER BY createdAt
DB_USERS --> MODEL : Array de usuarios
MODEL --> CTRL : Lista de usuarios
CTRL --> API : res.json(usuarios)
API --> PAGE : Array JSON
PAGE --> ACTOR : Renderiza tabla de usuarios
== Crear nuevo usuario ==
ACTOR -> PAGE : Completa formulario y hace clic Guardar
PAGE -> API : api.createUser({name, username, password, role, email})
API -> ROUTE : POST /api/users
ROUTE -> CTRL : userController.create(req, res)
CTRL -> MODEL : userModel.checkUsername(username)
MODEL -> DB_USERS : SELECT id FROM users WHERE username=?
alt Username ya existe
  DB_USERS --> MODEL : Fila encontrada
  MODEL --> CTRL : Username duplicado
  CTRL --> API : res.status(400) {error}
  API --> PAGE : Muestra toast de error
  PAGE --> ACTOR : Username ya registrado
else Username disponible
  DB_USERS --> MODEL : Sin resultado
  MODEL --> CTRL : Username disponible
  CTRL -> MODEL : userModel.create(userData)
  MODEL -> DB_USERS : INSERT INTO users VALUES(...)
  DB_USERS --> MODEL : insertId
  CTRL -> REP : registrarAccion(admin, "Creo usuario: "+username, "Usuarios", "")
  REP -> DB_HIST : INSERT INTO historial_acciones
  CTRL --> API : res.json({success, user})
  API --> PAGE : Usuario creado OK
  PAGE --> ACTOR : Muestra usuario en tabla
end
@enduml"""

PUML_SECUENCIA_CU011 = """\
@startuml Secuencia SAD CU011 MOPGIMED
title "Diagrama de Secuencia (Tecnico) — CU011: Visualizar Historial\\nMOPGIMED"
autonumber
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam sequenceArrowColor #B91C1C
skinparam sequenceArrowThickness 1.5
actor "Administrador" as ACTOR
box "Frontend (src/)" #EFF6FF
  participant "Usuarios.jsx" as PAGE
  participant "api.js" as API
end box
box "Backend (backend/)" #FFFBEB
  participant "routes/reportes.js" as ROUTE
  participant "reporteController.js" as CTRL
  participant "reporteModel.js" as MODEL
end box
box "Base de Datos" #ECFDF5
  participant "historial_acciones" as DB_HIST
  participant "users" as DB_USERS
end box
== Visualizar historial al cargar modulo ==
ACTOR -> PAGE : Navega a /usuarios
PAGE -> API : api.getHistorial()
API -> ROUTE : GET /api/reportes/historial
ROUTE -> CTRL : reporteController.getHistorial(req, res)
CTRL -> MODEL : reporteModel.getHistorial()
MODEL -> DB_HIST : SELECT * FROM historial_acciones ORDER BY fecha DESC
DB_HIST --> MODEL : Array de registros
MODEL --> CTRL : Lista de acciones
CTRL --> API : res.json(historial)
API --> PAGE : Array JSON
PAGE --> ACTOR : Renderiza tabla de historial en seccion inferior
== Registro automatico al gestionar usuarios ==
note over CTRL, DB_HIST : Este flujo es invocado internamente\\npor userController.js cada vez que\\nel Admin realiza una accion sobre usuarios
CTRL -> MODEL : reporteModel.registrarAccion(usuario, accion, modulo, detalle)
MODEL -> DB_USERS : SELECT name FROM users WHERE username=?
DB_USERS --> MODEL : Nombre del administrador
MODEL -> DB_HIST : INSERT INTO historial_acciones(fecha, usuario, accion, modulo, detalle)
DB_HIST --> MODEL : insertId confirmado
MODEL --> CTRL : Registro guardado en historial
@enduml"""

PUML_COMPONENTES = """\
@startuml MOPGIMED Diagrama de Componentes SAD
left to right direction
title "MOPGIMED — Diagrama de Componentes\\nDocumento de Arquitectura de Software (SAD)"
skinparam {
  backgroundColor #FFFFFF  shadowing false  defaultFontName Segoe UI
  ComponentBackgroundColor #FFFFFF  ComponentBorderColor #475569
  ArrowColor #B91C1C  ArrowThickness 1.5
  PackageBackgroundColor #F8FAFC  PackageBorderColor #94A3B8
  PackageFontStyle Bold  NodeBackgroundColor #1E293B
  NodeFontColor #FFFFFF  DatabaseBackgroundColor #ECFDF5
  DatabaseBorderColor #065F46
}
node "Navegador Web" as BROWSER
package "Frontend  —  localhost:5173" as FRONTEND #EFF6FF {
  component "React App\\n(src/main.jsx)" as REACT
  component "AuthContext.jsx" as CTX
  component "api.js (Axios)" as API
  component "Login.jsx" as PG_LOGIN
  component "Dashboard.jsx" as PG_DASH
  component "Usuarios.jsx" as PG_USR
  component "Medicamentos.jsx" as PG_MED
  component "ControlInventario.jsx" as PG_INV
  component "Notificaciones.jsx" as PG_NOTIF
  component "Analytics.jsx" as PG_ANA
  component "Reportes.jsx" as PG_REP
}
package "Backend  —  localhost:3001" as BACKEND #FFFBEB {
  component "Express App\\n(backend/index.js)" as EXPRESS
  component "routes/auth.js" as R_AUTH
  component "routes/users.js" as R_USR
  component "routes/medicamentos.js" as R_MED
  component "routes/inventario.js" as R_INV
  component "routes/notificaciones.js" as R_NOTIF
  component "routes/reportes.js" as R_REP
  component "routes/dashboard.js" as R_DASH
  component "routes/analytics.js" as R_ANA
  component "authController.js" as C_AUTH
  component "userController.js" as C_USR
  component "medicamentoController.js" as C_MED
  component "inventarioController.js" as C_INV
  component "notificacionController.js" as C_NOTIF
  component "reporteController.js" as C_REP
  component "dashboardController.js" as C_DASH
  component "analyticsController.js" as C_ANA
  component "userModel.js" as M_USR
  component "medicamentoModel.js" as M_MED
  component "reporteModel.js" as M_REP
  component "connection.js (MySQL2 Pool)" as POOL
}
database "MySQL\\nswipre_med" as MYSQL {
  component "users" as T_USR
  component "medicamentos" as T_MED
  component "historial_acciones" as T_HIST
}
BROWSER --> REACT : HTTP localhost:5173
REACT --> CTX : provee auth context
REACT --> API : llama a api.js
API --> EXPRESS : HTTP REST localhost:3001
EXPRESS --> R_AUTH
EXPRESS --> R_USR
EXPRESS --> R_MED
EXPRESS --> R_INV
EXPRESS --> R_NOTIF
EXPRESS --> R_REP
EXPRESS --> R_DASH
EXPRESS --> R_ANA
R_AUTH --> C_AUTH
R_USR  --> C_USR
R_MED  --> C_MED
R_INV  --> C_INV
R_NOTIF --> C_NOTIF
R_REP  --> C_REP
R_DASH --> C_DASH
R_ANA  --> C_ANA
C_AUTH --> M_USR
C_USR  --> M_USR
C_MED  --> M_MED
C_REP  --> M_REP
M_USR  --> POOL
M_MED  --> POOL
M_REP  --> POOL
POOL   --> MYSQL : MySQL2 Port 3306
@enduml"""

PUML_CLASES = """\
@startuml DiagramaClases MOPGIMED
title "Diagrama de Clases — MOPGIMED\\nSistema Inteligente de Inventario Farmacéutico"
top to bottom direction
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam defaultFontName Segoe UI
skinparam ArrowColor #B91C1C
skinparam ArrowThickness 1.5
skinparam class {
  BackgroundColor #FFFFFF
  BorderColor #64748B
  HeaderBackgroundColor #1E293B
  FontColor #EEEEEE
  AttributeFontColor #1E293B
  AttributeFontSize 9
}
skinparam package {
  BorderColor #94A3B8
  BackgroundColor #F8FAFC
  FontStyle Bold
  FontSize 12
}
enum RolUsuario { admin  farmacia  jefatura }
enum EstadoUsuario { Activo  Inactivo }
enum EstadoStock { Normal  Bajo  Critico }
package "Acceso" #EFF6FF {
  class Usuario {
    - id : int
    - username : String
    - password : String
    - role : RolUsuario
    - name : String
    - email : String
    - estado : EstadoUsuario
    - createdAt : Date
    ==
    + findByCredentials(u, p) : Usuario
    + findAll() : List<Usuario>
    + create(data) : Usuario
    + update(id, data) : Usuario
    + toggleStatus(id) : void
  }
}
package "Inventario" #ECFDF5 {
  class Medicamento {
    - id : int
    - codigo : String
    - nombre : String
    - laboratorio : String
    - lote : String
    - stock : int
    - stockMin : int
    - costoUnit : decimal
    - precioVenta : decimal
    - vencimiento : Date
    - categoria : String
    - estado : EstadoStock
    ==
    + findAll() : List<Medicamento>
    + create(data) : Medicamento
    + update(id, data) : Medicamento
    + remove(id) : void
    + calcEstado(stock, min) : EstadoStock
  }
}
package "Historial" #FFFBEB {
  class HistorialAccion {
    - id : int
    - fecha : String
    - usuario : String
    - accion : String
    - modulo : String
    - detalle : String
    ==
    + getHistorial() : List<HistorialAccion>
    + registrarAccion(u, a, m, d) : void
  }
}
package "Notificaciones" #FFF1F2 {
  class Notificacion {
    - id : int
    - tipo : String
    - titulo : String
    - descripcion : String
    - prioridad : String
    - leida : boolean
    ==
    + findAll() : List<Notificacion>
    + marcarLeida(id) : void
  }
}
Usuario "1" --> "0..*" HistorialAccion : registra >
Medicamento "1" --> "0..*" Notificacion : genera >
Usuario ..> RolUsuario
Usuario ..> EstadoUsuario
Medicamento ..> EstadoStock
@enduml"""


# ─────────────────────────────────────────────────────────────────────────────
#  DATOS TÉCNICOS
# ─────────────────────────────────────────────────────────────────────────────
RF_TECH = [
    ("UC-001", "Iniciar Sesión", "POST /api/auth/login", "authController.js → login()", "userModel.findByCredentials()"),
    ("UC-002", "Gestionar Usuarios", "GET /api/users\nPOST /api/users\nPUT /api/users/:id", "userController.js → getAll(), create(), update()", "userModel.findAll(), create(), update()"),
    ("UC-003", "Gestión Medicamentos CRUD", "GET /api/medicamentos\nPOST /api/medicamentos\nPUT /api/medicamentos/:id\nDELETE /api/medicamentos/:id", "medicamentoController.js → getAll(), create(), update(), remove()", "medicamentoModel.findAll(), create(), update(), remove()"),
    ("UC-004", "Búsqueda Medicamentos", "GET /api/medicamentos?search=&lab=&cat=", "medicamentoController.js → search()", "medicamentoModel.search()"),
    ("UC-005", "Control Inventario", "GET /api/inventario\nGET /api/inventario/vencimientos", "inventarioController.js → getInventario(), getVencimientos()", "inventarioModel.getInventario(), getVencimientos()"),
    ("UC-006", "Gestión Notificaciones", "GET /api/notificaciones\nPUT /api/notificaciones/:id\nPUT /api/notificaciones/all", "notificacionController.js → getAll(), marcarLeida(), marcarTodas()", "notificacionModel.findAll(), marcarLeida(), marcarTodas()"),
    ("UC-007", "Reporte General", "GET /api/reportes/resumen\nGET /api/reportes/categorias", "reporteController.js → getResumen(), getCategorias()", "reporteModel.getResumen(), getCategorias()"),
    ("UC-008", "Reporte Ingresos/Salidas", "GET /api/reportes/movimientos\nGET /api/reportes/comparativa", "reporteController.js → getMovimientos(), getComparativa()", "reporteModel.getMovimientos(), getComparativa()"),
    ("UC-009", "Reporte Rotación", "GET /api/reportes/rotacion", "reporteController.js → getRotacion()", "reporteModel.getRotacion()"),
    ("UC-010", "Dashboard Analítico", "GET /api/dashboard\nGET /api/analytics", "dashboardController.js, analyticsController.js", "dashboardModel, analyticsModel"),
    ("UC-011", "Historial Acciones", "GET /api/reportes/historial", "reporteController.js → getHistorial()", "reporteModel.getHistorial()"),
]

RNF_DATA = [
    ("RNF-001", "Seguridad", "Autenticación con credenciales únicas. Control de acceso basado en roles (RBAC). Contraseñas almacenadas con hash.", "Alta"),
    ("RNF-002", "Disponibilidad", "El sistema debe estar disponible durante el horario operativo del hospital con tiempo de inactividad máximo de 30 minutos por mantenimiento.", "Alta"),
    ("RNF-003", "Consistencia de datos", "Integridad referencial garantizada por MySQL. No se permiten registros huérfanos. Todas las escrituras son transaccionales.", "Alta"),
    ("RNF-004", "Rendimiento", "Carga de cualquier módulo en menos de 3 segundos. Consultas a la base de datos con tiempo de respuesta menor a 1 segundo.", "Media"),
    ("RNF-005", "Actualización automática", "El estado del stock y las notificaciones se recalculan automáticamente al guardar o editar un medicamento.", "Alta"),
    ("RNF-006", "Usabilidad", "Interfaz intuitiva basada en React 19. Semáforo de colores comprensible sin leyenda. Responsive design para escritorio.", "Media"),
    ("RNF-007", "Procesamiento analítico", "El módulo Analytics muestra predicciones de demanda con margen de error menor al 15% basadas en datos históricos del inventario.", "Media"),
]


# ─────────────────────────────────────────────────────────────────────────────
#  CONSTRUCCIÓN DEL DOCUMENTO SAD
# ─────────────────────────────────────────────────────────────────────────────
def build_sad():
    doc = Document()
    set_margins(doc, 2.5)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    def center(text, size=10, bold=False, color=C_NEGRO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.color.rgb = hex_to_rgb(color)
        return p

    # ── PORTADA ───────────────────────────────────────────────────────────────
    center(UNIV,     14, True,  C_OSCURO)
    center(FACULTAD, 12, False, C_OSCURO)
    center(ESCUELA,  11, False, C_OSCURO)
    doc.add_paragraph()
    center("FD04 — Documento de Arquitectura de Software", 16, True, C_ROJO)
    center("Sistema Inteligente de Inventario Farmacéutico", 14, True, C_OSCURO)
    center("MOPGIMED", 18, True, C_ROJO)
    doc.add_paragraph()
    center(f"Curso: {CURSO}", 11, False, C_OSCURO)
    center(f"Docente: {DOCENTE}", 11, False, C_OSCURO)
    doc.add_paragraph()
    center("Integrantes:", 11, True, C_OSCURO)
    for m in TEAM_MEMBERS:
        center(m, 10, False, C_OSCURO)
    doc.add_paragraph()
    center(f"Tacna, {YEAR}", 10, False, C_OSCURO)

    page_break(doc)

    # ── I. INTRODUCCIÓN ───────────────────────────────────────────────────────
    add_heading(doc, "I. INTRODUCCIÓN", 1)

    add_heading(doc, "1.1 Propósito", 2)
    add_body(doc, (
        "El presente Documento de Arquitectura de Software (SAD) describe la arquitectura del sistema "
        "MOPGIMED — Sistema Inteligente de Inventario Farmacéutico, desarrollado para la Clínica La Luz "
        "de Tacna. El documento define las decisiones arquitectónicas, las vistas estructurales del sistema "
        "y las relaciones entre sus componentes técnicos, siguiendo el modelo de vistas 4+1 de Kruchten."
    ))

    add_heading(doc, "1.2 Alcance", 2)
    add_body(doc, (
        "Este documento cubre la arquitectura completa del sistema MOPGIMED, incluyendo: el frontend React 19 "
        "con Vite, el backend Node.js con Express, la base de datos MySQL (swipre_med) y las interacciones "
        "entre todos sus componentes. El sistema está diseñado para operar en red local del hospital."
    ))

    add_heading(doc, "1.3 Definiciones y Acrónimos", 2)
    definitions = [
        ("MOPGIMED", "Sistema Inteligente de Inventario Farmacéutico — nombre del proyecto."),
        ("SRS", "Software Requirements Specification — Especificación de Requerimientos de Software."),
        ("SAD", "Software Architecture Document — Documento de Arquitectura de Software."),
        ("MVC", "Modelo-Vista-Controlador — patrón arquitectónico utilizado en el backend."),
        ("CRUD", "Create, Read, Update, Delete — operaciones básicas sobre la base de datos."),
        ("API", "Application Programming Interface — interfaz de comunicación entre frontend y backend."),
        ("REST", "Representational State Transfer — estilo arquitectónico de APIs HTTP."),
        ("JWT", "JSON Web Token — mecanismo de autenticación sin estado."),
        ("RBAC", "Role-Based Access Control — control de acceso basado en roles."),
        ("IA / ML", "Inteligencia Artificial / Machine Learning — módulo de predicción de demanda."),
        ("MySQL", "Sistema de gestión de base de datos relacional utilizado en el proyecto."),
        ("Vite", "Bundler de frontend de alto rendimiento utilizado con React 19."),
    ]
    def_table = doc.add_table(rows=1, cols=2)
    def_table.style = 'Table Grid'
    for i, hdr in enumerate(["Término", "Definición"]):
        set_bg(def_table.rows[0].cells[i], C_AZUL_OSCURO)
        cell_write(def_table.rows[0].cells[i], hdr, bold=True, size=9,
                   color=C_BLANCO, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, (term, defn) in enumerate(definitions):
        row = def_table.add_row()
        set_bg(row.cells[0], C_GRIS_CLARO if idx % 2 == 0 else C_BLANCO)
        set_bg(row.cells[1], C_BLANCO)
        cell_write(row.cells[0], term, bold=True, size=9)
        cell_write(row.cells[1], defn, size=9)
    doc.add_paragraph()

    add_heading(doc, "1.4 Referencias", 2)
    refs = [
        "FD03-SRS-MOPGIMED.docx — Especificación de Requerimientos de Software.",
        "Sprint_Backlog_MOPGIMED.docx — Registro de tareas de desarrollo.",
        "diagrama_casos_de_uso.puml — Diagrama de casos de uso del sistema.",
        "diagrama_clases.puml — Diagrama de clases del sistema.",
        "diagrama_paquetes_sad.puml — Diagrama de paquetes (vista técnica).",
        "diagrama_componentes_sad.puml — Diagrama de componentes del sistema.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ref)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

    add_heading(doc, "1.5 Visión General del Documento", 2)
    add_body(doc, (
        "El documento se organiza en 7 secciones: Introducción, Representación Arquitectónica, "
        "Objetivos y Limitaciones, Análisis de Requerimientos, Vistas de Caso de Uso, "
        "Diagrama de Componentes y Diagrama de Clases."
    ))

    page_break(doc)

    # ── II. REPRESENTACIÓN ARQUITECTÓNICA ─────────────────────────────────────
    add_heading(doc, "II. REPRESENTACIÓN ARQUITECTÓNICA", 1)

    add_heading(doc, "2.1 Escenarios", 2)
    escenarios = [
        ("Funcionalidad", "El sistema provee módulos para autenticación, gestión de medicamentos, control de inventario, notificaciones, reportes y análisis predictivo. Todos los módulos se comunican a través de una API REST."),
        ("Usabilidad", "La interfaz React 19 ofrece navegación intuitiva mediante un Sidebar fijo, validaciones en tiempo real, feedback visual con colores semáforo y modales para formularios."),
        ("Confiabilidad", "El backend Express valida todos los datos de entrada. La base de datos MySQL garantiza integridad referencial. El historial de acciones proporciona trazabilidad completa."),
        ("Rendimiento", "El frontend usa Vite para bundling optimizado. El backend utiliza un pool de conexiones MySQL2 para eficiencia en consultas concurrentes. Tiempo de respuesta objetivo: <3 segundos."),
        ("Mantenibilidad", "La arquitectura MVC separa claramente las responsabilidades. Los modelos encapsulan toda la lógica de base de datos. Los controladores son independientes entre sí."),
    ]
    for esc_name, esc_desc in escenarios:
        add_heading(doc, esc_name, 3)
        add_body(doc, esc_desc)

    add_heading(doc, "2.2 Vista Lógica — Diagrama de Paquetes (Técnico)", 2)
    add_body(doc, (
        "La vista lógica muestra la organización técnica del sistema en paquetes de archivos reales. "
        "El sistema se divide en tres bloques principales: Frontend (src/), Backend (backend/) y "
        "Base de Datos (MySQL: swipre_med). Las dependencias entre bloques siguen el patrón "
        "MVC: las rutas reciben las peticiones, los controladores procesan la lógica y los modelos "
        "acceden a la base de datos a través del pool de conexiones."
    ))
    add_code_block(doc, PUML_PAQUETES_SAD)
    add_caption(doc, "Figura 1: Diagrama de Paquetes (Vista Técnica SAD) — MOPGIMED")

    add_heading(doc, "2.3 Vista del Proceso", 2)
    add_body(doc, (
        "Los procesos principales del sistema son:"
    ))
    procesos = [
        "Proceso de autenticación: El usuario envía credenciales → Express valida → JWT generado → frontend almacena token → acceso a módulos según rol.",
        "Proceso de gestión de medicamentos: El personal completa formulario → Axios envía POST/PUT → Express valida → controlador llama al modelo → MySQL persiste → estado de stock calculado automáticamente.",
        "Proceso de notificaciones: Cada vez que se guarda un medicamento, el sistema evalúa el stock y genera notificaciones automáticas clasificadas por prioridad.",
        "Proceso de reportes: El usuario solicita un reporte → Express consulta MySQL → controlador agrega datos → respuesta JSON → frontend renderiza gráficas con los datos recibidos.",
        "Proceso de historial: Cada acción relevante (login, CRUD usuarios) invoca reporteController.registrarAccion() que inserta automáticamente en historial_acciones.",
    ]
    for proc in procesos:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(proc)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

    add_heading(doc, "2.4 Vista de Desarrollo", 2)
    tech_stack = [
        ("Frontend", "React 19", "Librería de interfaces de usuario basada en componentes"),
        ("Frontend", "Vite", "Bundler y servidor de desarrollo de alto rendimiento"),
        ("Frontend", "Axios", "Cliente HTTP para llamadas a la API REST"),
        ("Backend", "Node.js 20 LTS", "Entorno de ejecución JavaScript del lado del servidor"),
        ("Backend", "Express 4", "Framework web minimalista para Node.js"),
        ("Backend", "MySQL2", "Driver de MySQL con soporte para promesas y pool de conexiones"),
        ("Backend", "JWT (jsonwebtoken)", "Generación y verificación de tokens de autenticación"),
        ("Base de datos", "MySQL 8", "Sistema de gestión de base de datos relacional"),
        ("Herramientas", "python-docx", "Generación de documentos Word para reportes oficiales"),
        ("Herramientas", "PlantUML", "Generación de diagramas UML a partir de código"),
    ]
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    for i, hdr in enumerate(["Capa", "Tecnología", "Descripción"]):
        set_bg(tech_table.rows[0].cells[i], C_AZUL_OSCURO)
        cell_write(tech_table.rows[0].cells[i], hdr, bold=True, size=9,
                   color=C_BLANCO, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, (layer, tech, desc) in enumerate(tech_stack):
        row = tech_table.add_row()
        bg = C_GRIS_CLARO if idx % 2 == 0 else C_BLANCO
        for i, val in enumerate([layer, tech, desc]):
            set_bg(row.cells[i], bg)
            cell_write(row.cells[i], val, size=9)
    doc.add_paragraph()

    add_heading(doc, "2.5 Vista Física", 2)
    add_body(doc, (
        "El sistema MOPGIMED opera en la siguiente infraestructura física de red local:"
    ))
    add_body(doc, (
        "Servidor local → Express en puerto 3001 → MySQL en puerto 3306. "
        "El usuario accede desde su navegador a localhost:5173 (Vite dev server o build estático). "
        "La comunicación entre frontend y backend es HTTP REST en la misma red local del hospital. "
        "No se requiere conexión a internet para la operación normal del sistema."
    ))

    page_break(doc)

    # ── III. OBJETIVOS Y LIMITACIONES ─────────────────────────────────────────
    add_heading(doc, "III. OBJETIVOS Y LIMITACIONES ARQUITECTÓNICAS", 1)

    limitaciones = [
        ("Disponibilidad", "Alta", "El sistema debe estar disponible durante el horario operativo del hospital (07:00–22:00). Se diseña para operación en red local sin dependencia de servicios externos."),
        ("Seguridad", "Alta", "Control de acceso basado en roles (RBAC). Autenticación con JWT. Las contraseñas se almacenan con hash. Solo el Administrador puede gestionar usuarios y ver el historial."),
        ("Adaptabilidad", "Media", "La arquitectura MVC permite agregar nuevos módulos sin modificar los existentes. Los modelos encapsulan toda la lógica de base de datos para facilitar migraciones futuras."),
        ("Rendimiento", "Media", "Pool de conexiones MySQL2 para consultas concurrentes. Bundling optimizado con Vite. Lazy loading de componentes React para reducir el tiempo de carga inicial."),
        ("Escalabilidad", "Baja", "El sistema está diseñado para un único hospital. Para escalar a múltiples sedes se requeriría migrar a una base de datos en la nube y configurar un servidor de producción."),
    ]
    lim_table = doc.add_table(rows=1, cols=3)
    lim_table.style = 'Table Grid'
    for i, hdr in enumerate(["Atributo", "Prioridad", "Descripción"]):
        set_bg(lim_table.rows[0].cells[i], C_AZUL_OSCURO)
        cell_write(lim_table.rows[0].cells[i], hdr, bold=True, size=9,
                   color=C_BLANCO, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, (attr, prio, desc) in enumerate(limitaciones):
        row = lim_table.add_row()
        bg = C_GRIS_CLARO if idx % 2 == 0 else C_BLANCO
        for i, val in enumerate([attr, prio, desc]):
            set_bg(row.cells[i], bg)
            cell_write(row.cells[i], val, size=9)
    doc.add_paragraph()

    page_break(doc)

    # ── IV. ANÁLISIS DE REQUERIMIENTOS ────────────────────────────────────────
    add_heading(doc, "IV. ANÁLISIS DE REQUERIMIENTOS", 1)

    add_heading(doc, "4.1 Requerimientos Funcionales — Vista Técnica", 2)
    add_body(doc, "Mapeo de cada UC a su endpoint, controlador y modelo correspondiente:")

    rf_table = doc.add_table(rows=1, cols=5)
    rf_table.style = 'Table Grid'
    rf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rf_hdrs = ["UC", "Nombre", "Endpoint(s)", "Controlador → Método", "Modelo → Método"]
    rf_widths = [Cm(1.5), Cm(3.5), Cm(4.5), Cm(5.0), Cm(4.5)]
    for i, (hdr, w) in enumerate(zip(rf_hdrs, rf_widths)):
        cell = rf_table.rows[0].cells[i]
        cell.width = w
        set_bg(cell, C_AZUL_OSCURO)
        cell_write(cell, hdr, bold=True, size=8, color=C_BLANCO,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row_data in enumerate(RF_TECH):
        row = rf_table.add_row()
        bg = C_AZUL_CLARO if idx % 2 == 0 else C_BLANCO
        for i, val in enumerate(row_data):
            set_bg(row.cells[i], bg)
            cell_write(row.cells[i], val, size=8)
    doc.add_paragraph()

    add_heading(doc, "4.2 Requerimientos No Funcionales", 2)
    rnf_table = doc.add_table(rows=1, cols=4)
    rnf_table.style = 'Table Grid'
    rnf_hdrs = ["ID", "Nombre", "Descripción", "Prioridad"]
    rnf_widths = [Cm(1.8), Cm(3.5), Cm(10.0), Cm(2.0)]
    for i, (hdr, w) in enumerate(zip(rnf_hdrs, rnf_widths)):
        cell = rnf_table.rows[0].cells[i]
        cell.width = w
        set_bg(cell, C_AZUL_OSCURO)
        cell_write(cell, hdr, bold=True, size=9, color=C_BLANCO,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, (rnf_id, rnf_name, rnf_desc, rnf_prio) in enumerate(RNF_DATA):
        row = rnf_table.add_row()
        bg = C_VERDE_CLARO if idx % 2 == 0 else C_BLANCO
        for i, val in enumerate([rnf_id, rnf_name, rnf_desc, rnf_prio]):
            set_bg(row.cells[i], bg)
            cell_write(row.cells[i], val, size=9)
    doc.add_paragraph()

    page_break(doc)

    # ── V. VISTAS DE CASO DE USO ──────────────────────────────────────────────
    add_heading(doc, "V. VISTAS DE CASO DE USO", 1)

    add_heading(doc, "5.1 Diagrama de Casos de Uso", 2)
    add_code_block(doc, PUML_CASOS_USO)
    add_caption(doc, "Figura 2: Diagrama de Casos de Uso — MOPGIMED")

    add_heading(doc, "5.2 Diagramas de Robustez (Nombres Técnicos)", 2)
    add_body(doc, (
        "Los siguientes diagramas de robustez utilizan los nombres técnicos reales de los archivos del sistema "
        "(Boundary = archivos .jsx del frontend, Control = archivos Controller.js del backend, "
        "Entity = tablas de la base de datos y archivos Model.js)."
    ))

    fig_num = 3
    rob_data = [
        ("CU001", "Iniciar Sesión", PUML_ROB_SAD_CU001),
        ("CU002", "Gestionar Usuarios y Asignar Rol", PUML_ROB_SAD_CU002),
        ("CU011", "Visualizar Historial de Acciones", PUML_ROB_SAD_CU011),
    ]
    for rob_id, rob_name, rob_puml in rob_data:
        add_heading(doc, f"Robustez Técnica — {rob_id}: {rob_name}", 3)
        add_code_block(doc, rob_puml)
        add_caption(doc, f"Figura {fig_num}: Análisis de Robustez (SAD) — {rob_id}: {rob_name}")
        fig_num += 1

    add_heading(doc, "5.3 Diagramas de Secuencia (Nombres Técnicos)", 2)
    add_body(doc, (
        "Los diagramas de secuencia presentan el flujo técnico real de mensajes entre los componentes "
        "del sistema: páginas React, api.js, rutas Express, controladores, modelos y tablas MySQL."
    ))

    seq_data = [
        ("CU001", "Iniciar Sesión", PUML_SECUENCIA_CU001),
        ("CU002", "Gestionar Usuarios y Asignar Rol", PUML_SECUENCIA_CU002),
        ("CU011", "Visualizar Historial de Acciones", PUML_SECUENCIA_CU011),
    ]
    for seq_id, seq_name, seq_puml in seq_data:
        add_heading(doc, f"Secuencia Técnica — {seq_id}: {seq_name}", 3)
        add_code_block(doc, seq_puml)
        add_caption(doc, f"Figura {fig_num}: Diagrama de Secuencia (SAD) — {seq_id}: {seq_name}")
        fig_num += 1

    page_break(doc)

    # ── VI. DIAGRAMA DE COMPONENTES ───────────────────────────────────────────
    add_heading(doc, "VI. DIAGRAMA DE COMPONENTES", 1)
    add_body(doc, (
        "El diagrama de componentes muestra la arquitectura completa del sistema MOPGIMED, desde el "
        "navegador del usuario hasta la base de datos MySQL, pasando por todas las capas intermedias: "
        "páginas React, servicios Axios, rutas Express, controladores y modelos. Cada componente "
        "representa un archivo real del proyecto."
    ))
    add_code_block(doc, PUML_COMPONENTES)
    add_caption(doc, f"Figura {fig_num}: Diagrama de Componentes — MOPGIMED")
    fig_num += 1

    page_break(doc)

    # ── VII. DIAGRAMA DE CLASES ───────────────────────────────────────────────
    add_heading(doc, "VII. DIAGRAMA DE CLASES", 1)
    add_body(doc, (
        "El diagrama de clases muestra las entidades principales del sistema MOPGIMED, sus atributos, "
        "métodos y las relaciones entre ellas. Las clases corresponden a los modelos del backend y las "
        "tablas de la base de datos MySQL."
    ))
    add_code_block(doc, PUML_CLASES)
    add_caption(doc, f"Figura {fig_num}: Diagrama de Clases — MOPGIMED")

    # ── GUARDAR ───────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "FD04-SAD-MOPGIMED.docx")
    doc.save(output_path)
    print(f"[OK] Documento generado: {output_path}")


if __name__ == "__main__":
    build_sad()
